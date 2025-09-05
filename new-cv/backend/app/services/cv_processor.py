"""
CV Processing Service - Enhanced text extraction from PDF and DOCX files
"""
import logging
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Optional, Dict, Any
import time

# Suppress pdfminer warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdfinterp").setLevel(logging.ERROR)
logging.getLogger("pdfminer.layout").setLevel(logging.ERROR)

logger = logging.getLogger(__name__)


class CVProcessor:
    """Enhanced CV text extraction and processing"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
    
    def extract_text_from_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from CV file and return processing results
        
        Returns:
        {
            'text': str,
            'success': bool,
            'processing_time': float,
            'metadata': dict,
            'error': Optional[str]
        }
        """
        start_time = time.time()
        
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                result = self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                result = self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                result = self._extract_from_txt(file_path)
            else:
                return {
                    'text': '',
                    'success': False,
                    'processing_time': time.time() - start_time,
                    'metadata': {},
                    'error': f'Unsupported file format: {file_extension}'
                }
            
            processing_time = time.time() - start_time
            
            return {
                'text': result['text'],
                'success': True,
                'processing_time': processing_time,
                'metadata': {
                    'file_type': file_extension,
                    'text_length': len(result['text']),
                    'word_count': len(result['text'].split()) if result['text'] else 0,
                    **result.get('metadata', {})
                },
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error processing CV {file_path}: {str(e)}")
            return {
                'text': '',
                'success': False,
                'processing_time': time.time() - start_time,
                'metadata': {},
                'error': str(e)
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file using pdfplumber"""
        text_parts = []
        metadata = {'pages': 0}
        
        with pdfplumber.open(file_path) as pdf:
            metadata['pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up the text
                        page_text = self._clean_text(page_text)
                        text_parts.append(page_text)
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
        
        full_text = '\n\n'.join(text_parts)
        
        return {
            'text': full_text,
            'metadata': metadata
        }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract table content
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + tables_text
            full_text = '\n\n'.join(all_text)
            
            # Clean up the text
            full_text = self._clean_text(full_text)
            
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'has_tables': len(doc.tables) > 0
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    # Clean up the text
                    text = self._clean_text(text)
                    
                    return {
                        'text': text,
                        'metadata': {'encoding': encoding}
                    }
                    
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newlines
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Remove trailing spaces from lines
        lines = text.split('\n')
        cleaned_lines = [line.rstrip() for line in lines]
        text = '\n'.join(cleaned_lines)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """Get a preview of the text content"""
        if len(text) <= max_length:
            return text
        
        # Try to cut at a word boundary
        preview = text[:max_length]
        last_space = preview.rfind(' ')
        
        if last_space > max_length * 0.8:  # If we can cut at a word boundary reasonably close
            preview = preview[:last_space]
        
        return preview + "..."
    
    def extract_basic_info(self, text: str) -> Dict[str, Any]:
        """Extract basic information from CV text"""
        import re
        
        info = {
            'emails': [],
            'phones': [],
            'has_skills_section': False,
            'has_experience_section': False,
            'has_education_section': False
        }
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        info['emails'] = re.findall(email_pattern, text)
        
        # Extract phone numbers (basic pattern)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        info['phones'] = re.findall(phone_pattern, text)
        
        # Check for sections
        text_lower = text.lower()
        info['has_skills_section'] = any(keyword in text_lower for keyword in 
                                       ['skills', 'technical skills', 'competencies'])
        info['has_experience_section'] = any(keyword in text_lower for keyword in 
                                           ['experience', 'work experience', 'employment', 'career'])
        info['has_education_section'] = any(keyword in text_lower for keyword in 
                                          ['education', 'qualification', 'degree', 'university'])
        
        return info


# Global CV processor instance
cv_processor = CVProcessor()
