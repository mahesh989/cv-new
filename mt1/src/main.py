from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, PlainTextResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from uuid import uuid4
from datetime import datetime
from docx import Document
import os
import shutil
import json
import re
from dotenv import load_dotenv
from .job_tracker import router as job_tracker_router
from .ats_tester import router as ats_tester_router
from .cv_parser import extract_text_from_pdf, extract_text_from_docx
from .job_scraper import scrape_job_description
from .ai_matcher import analyze_match_fit
from .job_tracker import extract_metadata_gpt
# CV generation functionality moved to ats_tester.py
from .prompt_system import cv_tailoring_prompt, prompt_system
import openai
from .ai_config import get_model_params
from .ats_rules_engine import extract_skills_unified, evaluate_ats_compatibility
from .print_output_logger import append_output_log
from .job_queue_system import job_queue
# Utility imports for cleaner code organization
from .utils.data_formatters import (
    convert_to_frontend_format as _convert_to_frontend_format,
    consolidate_matched_skills as _consolidate_matched_skills,
    validate_comparison_completeness as _validate_comparison_completeness,
    parse_json_from_response,
    create_empty_skill_structure,
    extract_skill_list_from_text
)
from .utils.cv_parsers import (
    parse_cv_content_debug,
    process_section_content_debug
)

load_dotenv()

# Request models
class CVFilenameOnlyRequest(BaseModel):
    cv_filename: str

app = FastAPI()

# CORS setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories
UPLOAD_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "uploads"))
TAILORED_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "tailored_cvs"))
JOB_DB = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "job_db.json"))
ATS_DASHBOARD_DB = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "ats_dashboard.json"))

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TAILORED_DIR, exist_ok=True)

# Mount static files for serving tailored CVs
app.mount("/tailored_cvs", StaticFiles(directory=TAILORED_DIR), name="tailored_cvs")

@app.get("/")
def home():
    return {"status": "AI CV Agent is running!"}

@app.get("/health")
def health_check():
    """Health check endpoint for monitoring server status"""
    try:
        # Test basic imports
        import requests
        import bs4
        return {
            "status": "healthy",
            "server": "running",
            "dependencies": "ok",
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/ai-status/")
def get_ai_status():
    """Get the status of available AI providers."""
    from .hybrid_ai_service import hybrid_ai
    from .ai_config import model_state
    
    status = hybrid_ai.get_status()
    status['current_model'] = model_state.get_current_model()
    status['current_provider'] = model_state.get_current_provider()
    
    return status

@app.get("/list-cvs/")
def list_cvs():
    return {"uploaded_cvs": os.listdir(UPLOAD_DIR)}

@app.post("/upload-cv/")
async def upload_cv(cv: UploadFile = File(...)):
    path = os.path.join(UPLOAD_DIR, cv.filename)
    with open(path, "wb") as buffer:
        shutil.copyfileobj(cv.file, buffer)
    
    # 🔍 DEBUG: Print CV upload info
    print("=" * 80)
    print("🔍 DEBUG: CV UPLOAD INFO")
    print("=" * 80)
    print(f"CV Filename: {cv.filename}")
    print(f"CV File Path: {path}")
    print(f"CV File Size: {os.path.getsize(path)} bytes")
    print(f"CV Content Type: {cv.content_type}")
    print("=" * 80)
    
    return {"message": "CV uploaded.", "filename": cv.filename}

@app.post("/analyze-fit/")
async def analyze_fit(cv_filename: str = Form(...), text: str = Form(...)):
    path = os.path.join(UPLOAD_DIR, cv_filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="CV not found.")

    ext = os.path.splitext(cv_filename)[1].lower()
    if ext == ".pdf":
        cv_text = extract_text_from_pdf(path)
    elif ext == ".docx":
        cv_text = extract_text_from_docx(path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported CV format.")

    job_text = text.strip()
    if job_text.lower().startswith("http"):
        try:
            job_text = scrape_job_description(job_text)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Failed to scrape JD URL: {e}")
        if not job_text:
            raise HTTPException(status_code=422, detail="Empty job description.")

    try:
        # Use the same company name as CV/JD extraction for consistency
        company_name = "Maheshwor_Tiwari"
        
        
        result = await analyze_match_fit(cv_text, job_text, company_name)
        return JSONResponse(content=result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/scrape-job-description/")
async def scrape_jd(request: Request):
    data = await request.json()
    url = data.get("url")
    if not url:
        raise HTTPException(status_code=400, detail="No URL provided.")
    
    try:
        # Add timeout handling for the scraping operation
        import asyncio
        jd = await asyncio.wait_for(
            asyncio.to_thread(scrape_job_description, url),
            timeout=45.0  # 45 second timeout
        )
        
        # 🔍 DEBUG: Print JD text content
        print("=" * 80)
        print("🔍 DEBUG: SCRAPED JD TEXT CONTENT")
        print("=" * 80)
        print(f"URL: {url}")
        print(f"JD Text Length: {len(jd)} characters")
        print("JD Text Content:")
        print("-" * 40)
        print(jd[:1000] + "..." if len(jd) > 1000 else jd)
        print("-" * 40)
        print("=" * 80)
        
        return {"job_description": jd}
        
    except asyncio.TimeoutError:
        raise HTTPException(status_code=408, detail="Request timeout while scraping job description")
    except Exception as e:
        print(f"❌ Error in scrape_jd endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error scraping job description: {str(e)}")

@app.get("/get-prompt/")
def get_prompt():
    try:
        return {"prompt": prompt_system.get_prompt("cv_analysis", cv_text="{cv_text}", job_text="{job_text}")}
    except ImportError:
        return {"prompt": ""}

@app.post("/save-job/")
async def save_job(request: Request):
    try:
        data = await request.json()
        job_link = data.get("job_link")
        jd_text = data.get("jd_text")
        tailored_cv = data.get("tailored_cv")

        print(f"💾 Saving job with CV: {tailored_cv}")
        print(f"📁 Checking if file exists: {os.path.join(TAILORED_DIR, tailored_cv)}")

        if not job_link or not jd_text or not tailored_cv:
            raise HTTPException(status_code=400, detail="Missing required fields.")

        # Verify the CV file exists before saving
        cv_path = os.path.join(TAILORED_DIR, tailored_cv)
        if not os.path.exists(cv_path):
            print(f"❌ CV file not found before saving job: {cv_path}")
            raise HTTPException(status_code=404, detail="Tailored CV file not found.")

        # ✅ Use the GPT-powered metadata extraction
        metadata = extract_metadata_gpt(jd_text)

        job_data = {
            "sn": str(uuid4())[:8],
            "company": metadata.get("company", "Unknown"),
            "location": metadata.get("location", "Not specified"),
            "phone": metadata.get("phone", "Not found"),
            "date_applied": datetime.now().strftime("%Y-%m-%d"),
            "job_link": job_link,
            "tailored_cv": tailored_cv,
            "applied": bool(data.get("applied", False))  # default False if not given
        }

        print(f"📝 Saving job data: {job_data}")

        existing = []
        if os.path.exists(JOB_DB):
            with open(JOB_DB, "r") as f:
                try:
                    existing = json.load(f)
                except json.JSONDecodeError:
                    existing = []

        existing.append(job_data)
        with open(JOB_DB, "w") as f:
            json.dump(existing, f, indent=2)

        print(f"✅ Job saved successfully with CV: {tailored_cv}")
        return {"message": "Job saved successfully."}
    except Exception as e:
        print(f"❌ Error saving job: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error while saving job: {str(e)}")

@app.get("/jobs/")
def list_saved_jobs():
    if not os.path.exists(JOB_DB):
        return []
    with open(JOB_DB, "r") as f:
        try:
            return json.load(f)
        except json.JSONDecodeError:
            return []

@app.get("/list-jobs/")
def list_jobs_alias():
    """Alias for /jobs/ endpoint for compatibility"""
    return list_saved_jobs()

@app.post("/toggle-applied/")
async def toggle_applied(request: Request):
    try:
        data = await request.json()
        sn = data.get("sn")
        if not sn:
            raise HTTPException(status_code=400, detail="Missing job identifier.")

        if not os.path.exists(JOB_DB):
            raise HTTPException(status_code=404, detail="Job DB not found.")

        with open(JOB_DB, "r") as f:
            jobs = json.load(f)

        updated = False
        for job in jobs:
            if job.get("sn") == sn:
                job["applied"] = not job.get("applied", False)
                updated = True
                break

        if not updated:
            raise HTTPException(status_code=404, detail="Job not found.")

        with open(JOB_DB, "w") as f:
            json.dump(jobs, f, indent=2)

        return {"message": "Job application status toggled."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/tailored-cvs/{filename}", response_class=PlainTextResponse)
def get_cv_preview(filename: str):
    path = os.path.join(TAILORED_DIR, filename)
    print(f"🔍 Preview requested for: {path}")
    if not os.path.exists(path):
        print(f"❌ File not found: {path}")
        raise HTTPException(status_code=404, detail="File not found.")
    
    # Handle both PDF and DOCX files for preview
    if filename.lower().endswith('.pdf'):
        # Extract text from PDF for preview
        try:
            import PyPDF2
            with open(path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                print(f"✅ Successfully read PDF file: {path}")
                return text
        except Exception as e:
            # Fallback to pdfplumber if PyPDF2 fails
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
                    print(f"✅ Successfully read PDF file with pdfplumber: {path}")
                    return text
            except Exception as e2:
                print(f"❌ Error reading PDF file: {path}")
                print(f"Error details: {str(e2)}")
                raise HTTPException(status_code=500, detail=f"Failed to extract PDF text: {str(e2)}")
    else:
        # Handle DOCX files (backward compatibility)
        try:
            doc = Document(path)
            print(f"✅ Successfully read DOCX file: {path}")
            return "\n\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"❌ Error reading DOCX file: {path}")
            print(f"Error details: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to extract DOCX text: {str(e)}")

@app.get("/download-cv/{filename}")
def download_cv(filename: str):
    path = os.path.join(TAILORED_DIR, filename)
    print(f"📥 Download requested for: {path}")
    if not os.path.exists(path):
        print(f"❌ File not found: {path}")
        raise HTTPException(status_code=404, detail="CV file not found.")
    print(f"✅ File found, size: {os.path.getsize(path)} bytes")
    
    # Determine media type based on file extension
    if filename.lower().endswith('.pdf'):
        media_type = 'application/pdf'
    elif filename.lower().endswith('.docx'):
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        media_type = 'application/octet-stream'
    
    return FileResponse(
        path,
        filename=filename,
        media_type=media_type
    )

@app.get("/get-cv-content/{filename}", response_class=PlainTextResponse)
def get_cv_content(filename: str):
    """Get CV content for preview - works for both uploaded and tailored CVs"""
    # First try uploaded CVs directory (PDF/DOCX files)
    uploaded_path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(uploaded_path):
        try:
            # Use the existing CV parsing functions
            ext = os.path.splitext(filename)[1].lower()
            if ext == ".pdf":
                content = extract_text_from_pdf(uploaded_path)
            elif ext == ".docx":
                content = extract_text_from_docx(uploaded_path)
            else:
                # Try reading as text file for other formats
                with open(uploaded_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # 🔍 DEBUG: Print CV text content
            print("=" * 80)
            print("🔍 DEBUG: CV TEXT CONTENT (UPLOADED)")
            print("=" * 80)
            print(f"CV Filename: {filename}")
            print(f"CV Text Length: {len(content)} characters")
            print("CV Text Content:")
            print("-" * 40)
            print(content[:1000] + "..." if len(content) > 1000 else content)
            print("-" * 40)
            print("=" * 80)
            
            print(f"✅ Read uploaded CV: {filename} ({len(content)} chars)")
            return content
        except Exception as e:
            print(f"❌ Error reading uploaded CV {filename}: {str(e)}")
    
    # Then try tailored CVs directory (DOCX files)
    tailored_path = os.path.join(TAILORED_DIR, filename)
    if os.path.exists(tailored_path):
        try:
            doc = Document(tailored_path)
            content = "\n\n".join([para.text for para in doc.paragraphs])
            
            # 🔍 DEBUG: Print CV text content
            print("=" * 80)
            print("🔍 DEBUG: CV TEXT CONTENT (TAILORED)")
            print("=" * 80)
            print(f"CV Filename: {filename}")
            print(f"CV Text Length: {len(content)} characters")
            print("CV Text Content:")
            print("-" * 40)
            print(content[:1000] + "..." if len(content) > 1000 else content)
            print("-" * 40)
            print("=" * 80)
            
            print(f"✅ Read tailored CV: {filename} ({len(content)} chars)")
            return content
        except Exception as e:
            print(f"❌ Error reading tailored CV {filename}: {str(e)}")
    
    print(f"❌ CV file not found: {filename}")
    raise HTTPException(status_code=404, detail="CV file not found.")

@app.post("/print-cv-on-selection/")
async def print_cv_on_selection(cv_filename: str = Form(...)):
    """Print CV content exactly as it appears in the original file when CV is selected in the UI"""
    try:
        print("\n" + "="*80)
        print("🎯 CV SELECTED IN UI - PRINTING ORIGINAL CV CONTENT")
        print("="*80)
        print(f"📄 Selected CV: {cv_filename}")
        print(f"⏰ Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("="*80)
        
        # Get CV content directly from file without any processing
        uploaded_path = os.path.join(UPLOAD_DIR, cv_filename)
        if os.path.exists(uploaded_path):
            try:
                # Use the existing CV parsing functions to get raw content
                ext = os.path.splitext(cv_filename)[1].lower()
                if ext == ".pdf":
                    content = extract_text_from_pdf(uploaded_path)
                elif ext == ".docx":
                    content = extract_text_from_docx(uploaded_path)
                else:
                    # Try reading as text file for other formats
                    with open(uploaded_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                # Print the raw CV content exactly as extracted
                print("\n" + "="*80)
                print("📄 ORIGINAL CV CONTENT (EXACTLY AS IN FILE)")
                print("="*80)
                print(content)
                print("="*80)
                print("✅ Original CV content printed successfully")
                print("="*80 + "\n")
                
                return {"message": "Original CV content printed to console", "filename": cv_filename, "content_length": len(content)}
                
            except Exception as e:
                print(f"❌ Error reading uploaded CV {cv_filename}: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error reading CV file: {str(e)}")
        
        # Try tailored CVs directory if not found in uploads
        tailored_path = os.path.join(TAILORED_DIR, cv_filename)
        if os.path.exists(tailored_path):
            try:
                doc = Document(tailored_path)
                content = "\n\n".join([para.text for para in doc.paragraphs])
                
                # Print the raw CV content exactly as extracted
                print("\n" + "="*80)
                print("📄 ORIGINAL TAILORED CV CONTENT (EXACTLY AS IN FILE)")
                print("="*80)
                print(content)
                print("="*80)
                print("✅ Original tailored CV content printed successfully")
                print("="*80 + "\n")
                
                return {"message": "Original tailored CV content printed to console", "filename": cv_filename, "content_length": len(content)}
                
            except Exception as e:
                print(f"❌ Error reading tailored CV {cv_filename}: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error reading tailored CV file: {str(e)}")
        
        print(f"❌ CV file not found: {cv_filename}")
        raise HTTPException(status_code=404, detail="CV file not found.")
        
    except Exception as e:
        print(f"❌ Error printing CV {cv_filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error printing CV: {str(e)}")

@app.post("/print-cv-structured/")
async def print_cv_structured(cv_filename: str = Form(...)):
    """Print CV content in a structured format when CV is selected in the UI"""
    try:
        print("\n" + "="*80)
        print("🎯 CV SELECTED IN UI - STRUCTURED PRINT")
        print("="*80)
        print(f"📄 Selected CV: {cv_filename}")
        print(f"⏰ Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("="*80)
        
        # Get CV content
        content = get_cv_content(cv_filename)
        
        # Parse CV using Claude for structured output
        from .cv_accuracy_enhancer import parse_resume_with_claude
        parsed_data = parse_resume_with_claude(content)
        
        # Print structured CV information
        print("\n" + "="*80)
        print("📄 STRUCTURED CV INFORMATION")
        print("="*80)
        
        # Contact Information
        if 'contact_information' in parsed_data:
            contact = parsed_data['contact_information']
            print("📧 CONTACT INFORMATION")
            print("-" * 40)
            if 'name' in contact:
                print(f"Name: {contact['name']}")
            if 'email' in contact:
                print(f"Email: {contact['email']}")
            if 'phone' in contact:
                print(f"Phone: {contact['phone']}")
            print()
        
        # Summary
        if 'summary' in parsed_data:
            print("📝 SUMMARY")
            print("-" * 40)
            print(parsed_data['summary'])
            print()
        
        # Experience
        if 'experience' in parsed_data:
            print("💼 EXPERIENCE")
            print("-" * 40)
            for i, exp in enumerate(parsed_data['experience'], 1):
                print(f"{i}. {exp.get('title', 'N/A')} at {exp.get('company', 'N/A')}")
                print(f"   {exp.get('duration', 'N/A')}")
                print(f"   {exp.get('location', 'N/A')}")
                if 'bullets' in exp and exp['bullets']:
                    print("   Bullets:")
                    for bullet in exp['bullets'][:3]:  # Show first 3 bullets
                        print(f"   • {bullet}")
                print()
        
        # Education
        if 'education' in parsed_data:
            print("🎓 EDUCATION")
            print("-" * 40)
            for i, edu in enumerate(parsed_data['education'], 1):
                print(f"{i}. {edu.get('degree', 'N/A')}")
                print(f"   {edu.get('institution', 'N/A')}")
                print(f"   {edu.get('duration', 'N/A')}")
                print()
        
        # Skills
        if 'skills' in parsed_data:
            print("🔧 SKILLS")
            print("-" * 40)
            skills = parsed_data['skills']
            if isinstance(skills, list):
                print(", ".join(skills))
            elif isinstance(skills, dict):
                for category, skill_list in skills.items():
                    if skill_list:
                        print(f"{category}: {', '.join(skill_list)}")
            print()
        
        # Projects
        if 'projects' in parsed_data:
            print("🚀 PROJECTS")
            print("-" * 40)
            for i, project in enumerate(parsed_data['projects'], 1):
                print(f"{i}. {project.get('title', 'N/A')}")
                print(f"   {project.get('description', 'N/A')}")
                if 'technologies' in project and project['technologies']:
                    print(f"   Technologies: {', '.join(project['technologies'])}")
                print()
        
        # Languages
        if 'languages' in parsed_data:
            print("🌍 LANGUAGES")
            print("-" * 40)
            languages = parsed_data['languages']
            if isinstance(languages, list):
                print(", ".join(languages))
            print()
        
        print("="*80)
        print("✅ Structured CV information printed successfully")
        print("="*80 + "\n")
        
        return {
            "message": "Structured CV information printed to console", 
            "filename": cv_filename, 
            "content_length": len(content),
            "parsed_sections": list(parsed_data.keys())
        }
        
    except Exception as e:
        print(f"❌ Error printing structured CV {cv_filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error printing structured CV: {str(e)}")

@app.post("/print-cv-original/")
async def print_cv_original(cv_filename: str = Form(...)):
    """Print CV content exactly as it appears in the original file when CV is selected in the UI"""
    try:
        # Get CV content directly from file without any processing
        uploaded_path = os.path.join(UPLOAD_DIR, cv_filename)
        if os.path.exists(uploaded_path):
            try:
                # Use the existing CV parsing functions to get raw content
                ext = os.path.splitext(cv_filename)[1].lower()
                if ext == ".pdf":
                    content = extract_text_from_pdf(uploaded_path)
                elif ext == ".docx":
                    content = extract_text_from_docx(uploaded_path)
                else:
                    # Try reading as text file for other formats
                    with open(uploaded_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                # Print beautifully formatted CV content
                print("\n" + "🎯" + "="*78 + "🎯")
                print("📄 CV SELECTED - PRINTING ORIGINAL CONTENT")
                print("🎯" + "="*78 + "🎯")
                print(f"📋 Filename: {cv_filename}")
                print(f"⏰ Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                print(f"📏 Content Length: {len(content)} characters")
                print("🎯" + "="*78 + "🎯")
                print()
                
                # Print the CV content with proper formatting
                print("📄 CV CONTENT:")
                print("─" * 80)
                print(content)
                print("─" * 80)
                print()
                print("✅ CV content printed successfully!")
                print("🎯" + "="*78 + "🎯\n")
                
                return {"message": "Original CV content printed to console", "filename": cv_filename, "content_length": len(content)}
                
            except Exception as e:
                print(f"❌ Error reading uploaded CV {cv_filename}: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error reading CV file: {str(e)}")
        
        # Try tailored CVs directory if not found in uploads
        tailored_path = os.path.join(TAILORED_DIR, cv_filename)
        if os.path.exists(tailored_path):
            try:
                doc = Document(tailored_path)
                content = "\n\n".join([para.text for para in doc.paragraphs])
                
                # Print beautifully formatted tailored CV content
                print("\n" + "🎯" + "="*78 + "🎯")
                print("📄 TAILORED CV SELECTED - PRINTING ORIGINAL CONTENT")
                print("🎯" + "="*78 + "🎯")
                print(f"📋 Filename: {cv_filename}")
                print(f"⏰ Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                print(f"📏 Content Length: {len(content)} characters")
                print("🎯" + "="*78 + "🎯")
                print()
                
                # Print the CV content with proper formatting
                print("📄 TAILORED CV CONTENT:")
                print("─" * 80)
                print(content)
                print("─" * 80)
                print()
                print("✅ Tailored CV content printed successfully!")
                print("🎯" + "="*78 + "🎯\n")
                
                return {"message": "Original tailored CV content printed to console", "filename": cv_filename, "content_length": len(content)}
                
            except Exception as e:
                print(f"❌ Error reading tailored CV {cv_filename}: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error reading tailored CV file: {str(e)}")
        
        print(f"❌ CV file not found: {cv_filename}")
        raise HTTPException(status_code=404, detail="CV file not found.")
        
    except Exception as e:
        print(f"❌ Error printing CV {cv_filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error printing CV: {str(e)}")

@app.get("/debug-cv-parsing/{filename}")
def debug_cv_parsing(filename: str):
    """Debug endpoint to print parsed CV JSON structure"""
    try:
        # Get CV content
        uploaded_path = os.path.join(UPLOAD_DIR, filename)
        if os.path.exists(uploaded_path):
            ext = os.path.splitext(filename)[1].lower()
            if ext == ".pdf":
                content = extract_text_from_pdf(uploaded_path)
            elif ext == ".docx":
                content = extract_text_from_docx(uploaded_path)
            else:
                with open(uploaded_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            print("=" * 80)
            print("🔍 [DEBUG] CV PARSING DEBUG OUTPUT")
            print("=" * 80)
            print(f"📄 Filename: {filename}")
            print(f"📏 Content Length: {len(content)} characters")
            print("=" * 80)
            print("📄 RAW CV CONTENT (first 1000 chars):")
            print("=" * 80)
            print(content[:1000])
            print("=" * 80)
            
            # Parse the content using the same logic as frontend
            parsed_sections = parse_cv_content_debug(content)
            
            print("📊 PARSED SECTIONS:")
            print("=" * 80)
            for i, section in enumerate(parsed_sections):
                print(f"Section {i+1}: {section.get('section_title', 'Unknown')}")
                print(f"  Content items: {len(section.get('content', []))}")
                for j, item in enumerate(section.get('content', [])[:3]):  # Show first 3 items
                    item_type = item.get('type', 'unknown')
                    item_text = item.get('text', '')[:100]  # Truncate long text
                    print(f"    Item {j+1} ({item_type}): {item_text}")
                if len(section.get('content', [])) > 3:
                    print(f"    ... and {len(section.get('content', [])) - 3} more items")
                print()
            
            return {
                "filename": filename,
                "content_length": len(content),
                "parsed_sections": parsed_sections,
                "total_sections": len(parsed_sections)
            }
            
        else:
            print(f"❌ CV file not found: {filename}")
            return {"error": "CV file not found"}
            
    except Exception as e:
        print(f"❌ Error in debug parsing: {str(e)}")
        return {"error": str(e)}

def parse_cv_content_debug(raw_content):
    """Debug version of CV parsing logic"""
    sections = []
    lines = raw_content.split('\n')
    
    current_section = ''
    current_content = []
    
    # Major CV section headers (standalone, all caps)
    section_headers = [
        'EDUCATION',
        'EXPERIENCE',
        'WORK EXPERIENCE',
        'EMPLOYMENT HISTORY',
        'PROJECTS',
        'SKILLS',
        'TECHNICAL SKILLS',
        'CERTIFICATIONS',
        'AWARDS',
        'PUBLICATIONS',
        'VOLUNTEER',
        'INTERESTS',
        'REFERENCES'
    ]

    for i, line in enumerate(lines):
        trimmed_line = line.strip()
        if trimmed_line == '':
            continue

        # Check if this is a major section header (standalone, all caps)
        is_header = False
        matched_header = ''
        
        # Only treat as header if it's a standalone line in all caps
        if trimmed_line == trimmed_line.upper() and len(trimmed_line) > 2:
            for header in section_headers:
                if trimmed_line == header.upper():
                    is_header = True
                    matched_header = header
                    break

        if is_header:
            # Save previous section if exists
            if current_section and current_content:
                sections.append({
                    'section_title': current_section,
                    'content': process_section_content_debug(current_content, current_section),
                })
            
            current_section = matched_header if matched_header else trimmed_line
            current_content = []
        else:
            # Handle contact information and personal info
            if not sections and not current_section:
                if ('@' in trimmed_line or 
                    '|' in trimmed_line or
                    'Phone' in trimmed_line or
                    'LinkedIn' in trimmed_line or
                    'GitHub' in trimmed_line or
                    'Blogs' in trimmed_line or
                    'Portfolio' in trimmed_line):
                    if not current_content:
                        current_section = 'CONTACT INFORMATION'
                elif i < 5 and not current_content:
                    current_section = 'PERSONAL INFORMATION'
            
            # Handle career profile section (special case)
            if trimmed_line.upper() == 'CAREER PROFILE':
                if current_section and current_content:
                    sections.append({
                        'section_title': current_section,
                        'content': process_section_content_debug(current_content, current_section),
                    })
                current_section = 'CAREER PROFILE'
                current_content = []
            else:
                current_content.append(trimmed_line)

    # Add the last section
    if current_section and current_content:
        sections.append({
            'section_title': current_section,
            'content': process_section_content_debug(current_content, current_section),
        })

    # If no sections were found, create a general content section
    if not sections:
        sections.append({
            'section_title': 'CV Content',
            'content': process_section_content_debug([line for line in lines if line.strip()]),
        })

    return sections

def process_section_content_debug(content, section_title=None):
    """Debug version of content processing with lookahead for job titles in EXPERIENCE section"""
    processed_content = []
    i = 0
    while i < len(content):
        trimmed_line = content[i].strip()
        if trimmed_line == '':
            i += 1
            continue

        # Check if line contains bullet points
        if (trimmed_line.startswith('•') or
            trimmed_line.startswith('*') or
            trimmed_line.startswith('-') or
            trimmed_line.startswith('▪') or
            trimmed_line.startswith('▫') or
            re.match(r'^\d+\.', trimmed_line)):
            # Extract bullet content
            bullet_content = trimmed_line
            if trimmed_line.startswith('•'):
                bullet_content = trimmed_line[1:].strip()
            elif trimmed_line.startswith('*'):
                bullet_content = trimmed_line[1:].strip()
            elif trimmed_line.startswith('-'):
                bullet_content = trimmed_line[1:].strip()
            elif trimmed_line.startswith('▪'):
                bullet_content = trimmed_line[1:].strip()
            elif trimmed_line.startswith('▫'):
                bullet_content = trimmed_line[1:].strip()
            elif re.match(r'^\d+\.', trimmed_line):
                bullet_content = re.sub(r'^\d+\.\s*', '', trimmed_line)
            processed_content.append({
                'type': 'bullet',
                'text': bullet_content,
            })
            i += 1
            continue

        # Check if this is a job title/position line (contains date range)
        is_job_title = False
        if (re.search(r'\d{4}\s*[-–—]\s*(Present|\d{4})', trimmed_line) or
            re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', trimmed_line)):
            is_job_title = True

        # Improved: If in EXPERIENCE section, treat as job title if next line is a date
        if (section_title is not None and 'EXPERIENCE' in section_title.upper()) and not is_job_title:
            if i + 1 < len(content):
                next_line = content[i + 1].strip()
                if (re.search(r'\d{4}\s*[-–—]\s*(Present|\d{4})', next_line) or
                    re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', next_line)):
                    is_job_title = True
        if is_job_title:
            processed_content.append({
                'type': 'job_title',
                'text': trimmed_line,
            })
            i += 1
            continue

        # Education institution
        if (',' in trimmed_line and 
            ('University' in trimmed_line or 
             'College' in trimmed_line or
             'School' in trimmed_line)):
            processed_content.append({
                'type': 'education',
                'text': trimmed_line,
            })
            i += 1
            continue

        # Regular text line
        processed_content.append({
            'type': 'text',
            'text': trimmed_line,
        })
        i += 1
    return processed_content

@app.delete("/delete-job/{sn}")
def delete_job(sn: str):
    if not os.path.exists(JOB_DB):
        raise HTTPException(status_code=404, detail="Job DB not found.")

    with open(JOB_DB, "r") as f:
        jobs = json.load(f)

    updated_jobs = [job for job in jobs if job.get("sn") != sn]

    with open(JOB_DB, "w") as f:
        json.dump(updated_jobs, f, indent=2)

    return {"message": "Job deleted successfully."}

@app.post("/extract-metadata/")
async def extract_metadata(request: Request):
    data = await request.json()
    jd_text = data.get("jd_text")
    if not jd_text:
        raise HTTPException(status_code=400, detail="Missing job description text.")
    metadata = extract_metadata_gpt(jd_text)
    return metadata

@app.post("/api/llm/extract-job-info")
async def extract_job_info(request: Request):
    """Extract job title and company name from job description using LLM"""
    try:
        data = await request.json()
        job_description = data.get("job_description", "")
        
        if not job_description:
            raise HTTPException(status_code=400, detail="No job description provided")
        
        try:
            # Use Claude API for consistency with other endpoints
            from .hybrid_ai_service import hybrid_ai
            
            # Create a prompt for extracting job information
            prompt = f"""
            Analyze the following job description and extract the job title and company name.
            
            CRITICAL RULES:
            1. Extract ONLY information that ACTUALLY APPEARS in the text
            2. Do NOT infer or assume information not explicitly stated
            3. Return result in JSON format with exactly these keys: "job_title" and "company"
            4. If you cannot find either piece of information, return null for that field
            5. Be precise and extract only the actual job title and company name, not additional descriptive text
            
            Job Description:
            {job_description}
            
            Return ONLY a JSON object in this exact format:
            {{"job_title": "actual job title or null", "company": "actual company name or null"}}
            """
            
            # Use DeepSeek API
            result_text = await hybrid_ai.generate_response(
                prompt=prompt,
                temperature=0.0,
                max_tokens=200
            )
            
            print(f"🔍 [JOB-INFO-CLAUDE] Raw response: {result_text}")
            
            # Try to extract JSON from the response
            try:
                # Find JSON in the response
                import re
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    result_json = json.loads(json_match.group(0))
                else:
                    result_json = json.loads(result_text)
                
                # Ensure we have the required keys
                job_title = result_json.get("job_title")
                company = result_json.get("company")
                
                return {
                    "job_title": job_title,
                    "company": company
                }
                
            except json.JSONDecodeError:
                # Fallback: try to parse manually
                lines = result_text.split('\n')
                job_title = None
                company = None
                
                for line in lines:
                    if 'job_title' in line.lower() and ':' in line:
                        job_title = line.split(':', 1)[1].strip().strip('"').strip("'")
                    elif 'company' in line.lower() and ':' in line:
                        company = line.split(':', 1)[1].strip().strip('"').strip("'")
                
                return {
                    "job_title": job_title,
                    "company": company
                }
                
        except Exception as claude_error:
            print(f"❌ [JOB-INFO-CLAUDE] Error: {claude_error}, falling back to GPT")
            
            # Fallback to original GPT method
            prompt = f"""
            Analyze the following job description and extract the job title and company name.
            Return the result in JSON format with exactly these keys: "job_title" and "company".
            If you cannot find either piece of information, return null for that field.
            Be precise and extract only the actual job title and company name, not additional descriptive text.
            
            Job Description:
            {job_description}
            
            JSON Response:
            """
            
            # Use OpenAI to extract information
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a job information extraction assistant. Always respond with valid JSON containing job_title and company fields."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=200
            )
            
            # Parse the response
            result_text = response.choices[0].message.content.strip()
            
            # Try to extract JSON from the response
            try:
                # Find JSON in the response
                import re
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    result_json = json.loads(json_match.group(0))
                else:
                    result_json = json.loads(result_text)
                
                # Ensure we have the required keys
                job_title = result_json.get("job_title")
                company = result_json.get("company")
                
                return {
                    "job_title": job_title,
                    "company": company
                }
                
            except json.JSONDecodeError:
                # Fallback: try to parse manually
                lines = result_text.split('\n')
                job_title = None
                company = None
                
                for line in lines:
                    if 'job_title' in line.lower() and ':' in line:
                        job_title = line.split(':', 1)[1].strip().strip('"').strip("'")
                    elif 'company' in line.lower() and ':' in line:
                        company = line.split(':', 1)[1].strip().strip('"').strip("'")
                
                return {
                    "job_title": job_title,
                    "company": company
                }
            
    except Exception as e:
        print(f"❌ Error extracting job info: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting job information: {str(e)}")

# ================================
# ATS DASHBOARD ENDPOINTS  
# ================================

@app.post("/ats-dashboard/save-result/")
async def save_ats_result(request: Request):
    """Save ATS test result to dashboard"""
    try:
        data = await request.json()
        
        # Validate required fields
        required_fields = ["jobId", "jobTitle", "company", "atsScore", "cvName", "matchedSkills", "missedSkills"]
        for field in required_fields:
            if field not in data:
                raise HTTPException(status_code=400, detail=f"Missing required field: {field}")
        
        ats_result = {
            "jobId": data["jobId"],
            "jobTitle": data["jobTitle"],
            "company": data["company"],
            "jdText": data.get("jdText", ""),
            "testDate": data.get("testDate", datetime.now().isoformat()),
            "atsScore": data["atsScore"],
            "cvName": data["cvName"],
            "matchedSkills": data["matchedSkills"],
            "missedSkills": data["missedSkills"],
            "metadata": data.get("metadata", {}),
            "status": data.get("status", "completed"),
            "createdAt": datetime.now().isoformat(),
            "updatedAt": datetime.now().isoformat()
        }
        
        # Load existing data
        existing_results = []
        if os.path.exists(ATS_DASHBOARD_DB):
            with open(ATS_DASHBOARD_DB, "r") as f:
                try:
                    existing_results = json.load(f)
                except json.JSONDecodeError:
                    existing_results = []
        
        # Check if job already exists and update or add
        existing_index = -1
        for i, result in enumerate(existing_results):
            if result.get("jobId") == ats_result["jobId"]:
                existing_index = i
                break
        
        if existing_index != -1:
            # Update existing result - preserve test history
            existing_result = existing_results[existing_index]
            existing_metadata = existing_result.get("metadata", {})
            test_history = existing_metadata.get("testHistory", [])
            
            # Add current test to history
            test_history.append({
                "date": ats_result["testDate"],
                "score": ats_result["atsScore"],
                "matchedSkills": ats_result["matchedSkills"],
                "missedSkills": ats_result["missedSkills"],
                "cvName": ats_result["cvName"]
            })
            
            # Update metadata
            ats_result["metadata"]["testHistory"] = test_history
            ats_result["metadata"]["totalTests"] = len(test_history)
            
            existing_results[existing_index] = ats_result
            print(f"✅ Updated ATS result: {ats_result['jobTitle']} at {ats_result['company']} ({len(test_history)} tests)")
        else:
            # Add new result
            existing_results.append(ats_result)
            print(f"✅ Added new ATS result: {ats_result['jobTitle']} at {ats_result['company']}")
        
        # Save to file
        with open(ATS_DASHBOARD_DB, "w") as f:
            json.dump(existing_results, f, indent=2)
        
        return {"message": "ATS result saved successfully"}
        
    except Exception as e:
        print(f"❌ Error saving ATS result: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving ATS result: {str(e)}")

@app.get("/ats-dashboard/results/")
def get_ats_results():
    """Get all ATS dashboard results"""
    try:
        if not os.path.exists(ATS_DASHBOARD_DB):
            print("📭 No ATS dashboard file found")
            return []
        
        with open(ATS_DASHBOARD_DB, "r") as f:
            try:
                results = json.load(f)
                # Sort by test date (newest first)
                results.sort(key=lambda x: x.get("testDate", ""), reverse=True)
                print(f"📊 Loaded {len(results)} ATS results")
                return results
            except json.JSONDecodeError:
                print("❌ Invalid JSON in ATS dashboard file")
                return []
                
    except Exception as e:
        print(f"❌ Error loading ATS results: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error loading ATS results: {str(e)}")

@app.delete("/ats-dashboard/clear/")
def clear_ats_results():
    """Clear all ATS dashboard results"""
    try:
        if os.path.exists(ATS_DASHBOARD_DB):
            # Keep the file but make it empty array
            with open(ATS_DASHBOARD_DB, "w") as f:
                json.dump([], f)
            print("🗑️ Cleared all ATS dashboard results")
        
        return {"message": "All ATS results cleared successfully"}
        
    except Exception as e:
        print(f"❌ Error clearing ATS results: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error clearing ATS results: {str(e)}")

@app.get("/ats-dashboard/stats/")
def get_ats_stats():
    """Get ATS dashboard statistics"""
    try:
        if not os.path.exists(ATS_DASHBOARD_DB):
            return {
                "totalJobs": 0,
                "avgScore": 0,
                "topScore": 0,
                "storageType": "JSON File"
            }
        
        with open(ATS_DASHBOARD_DB, "r") as f:
            try:
                results = json.load(f)
                
                if not results:
                    return {
                        "totalJobs": 0,
                        "avgScore": 0,
                        "topScore": 0,
                        "storageType": "JSON File"
                    }
                
                scores = [result.get("atsScore", 0) for result in results]
                avg_score = sum(scores) / len(scores) if scores else 0
                top_score = max(scores) if scores else 0
                
                return {
                    "totalJobs": len(results),
                    "avgScore": round(avg_score),
                    "topScore": top_score,
                    "storageType": "JSON File"
                }
                
            except json.JSONDecodeError:
                return {
                    "totalJobs": 0,
                    "avgScore": 0,
                    "topScore": 0,
                    "storageType": "JSON File"
                }
                
    except Exception as e:
        print(f"❌ Error getting ATS stats: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting ATS stats: {str(e)}")

@app.delete("/ats-dashboard/delete/{job_id}")
def delete_ats_result(job_id: str):
    """Delete a specific ATS result by job ID"""
    try:
        if not os.path.exists(ATS_DASHBOARD_DB):
            raise HTTPException(status_code=404, detail="ATS dashboard not found")
        
        with open(ATS_DASHBOARD_DB, "r") as f:
            try:
                results = json.load(f)
            except json.JSONDecodeError:
                results = []
        
        # Find and remove the result with matching jobId
        original_count = len(results)
        results = [result for result in results if result.get("jobId") != job_id]
        
        if len(results) == original_count:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Save the updated results
        with open(ATS_DASHBOARD_DB, "w") as f:
            json.dump(results, f, indent=2)
        
        print(f"🗑️ Deleted ATS result with job ID: {job_id}")
        return {"message": "ATS result deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error deleting ATS result: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting ATS result: {str(e)}")

# Duplicate DOCX endpoints removed - main endpoints already handle DOCX format

# CV generation endpoints are now in ats_tester_router

app.include_router(job_tracker_router)

app.include_router(ats_tester_router)

@app.post("/parse-resume/")
async def parse_resume(file: UploadFile = File(...)):
    import tempfile, os, json
    # Save uploaded file to a temp location
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        # Extract text
        def extract_text(file_path):
            import docx2txt
            import fitz
            if file_path.endswith('.pdf'):
                doc = fitz.open(file_path)
                return "\n".join([page.get_text() for page in doc])
            elif file_path.endswith('.docx'):
                return docx2txt.process(file_path)
            else:
                raise ValueError("Unsupported file format")
        resume_text = extract_text(tmp_path)
        
        # Parse using Claude
        from .cv_accuracy_enhancer import parse_resume_with_claude
        parsed_data = parse_resume_with_claude(resume_text)
        
        return JSONResponse(content=parsed_data)
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

@app.post("/parse-uploaded-cv/")
async def parse_uploaded_cv(cv_filename: str = Form(...)):
    """Parse an already uploaded CV file by filename"""
    try:
        # Check if file exists in uploads directory
        file_path = os.path.join(UPLOAD_DIR, cv_filename)
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="CV file not found")
        
        # Extract text based on file extension
        ext = os.path.splitext(cv_filename)[1].lower()
        if ext == ".pdf":
            resume_text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            resume_text = extract_text_from_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        # Parse using Claude
        from .cv_accuracy_enhancer import parse_resume_with_claude
        parsed_data = parse_resume_with_claude(resume_text)
        
        return JSONResponse(content=parsed_data)
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/test-unified-extractor/")
async def test_unified_extractor_endpoint(payload: CVFilenameOnlyRequest):
    """Test the unified skill extractor on a CV file"""
    cv_filename = payload.cv_filename
    try:
        file_path = os.path.join(UPLOAD_DIR, cv_filename)
        if not os.path.exists(file_path):
            return JSONResponse(
                status_code=404,
                content={"error": "CV not found."}
            )
        
        ext = os.path.splitext(cv_filename)[1].lower()
        if ext == ".pdf":
            cv_text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            cv_text = extract_text_from_docx(file_path)
        else:
            return JSONResponse(
                status_code=400,
                content={"error": "Unsupported CV format."}
            )

        # Extract skills using unified extractor
        print(f"\n🔍 SKILL EXTRACTION ENDPOINT CALLED")
        print("="*50)
        print(f"📄 CV Filename: {cv_filename}")
        print(f"📝 CV Length: {len(cv_text)} characters")
        print(f"📄 CV Text (first 300 chars): {cv_text[:300]}...")
        print("="*50)
        
        skills_result = await extract_skills_unified(cv_text)
        
        print(f"\n✅ SKILL EXTRACTION RESULTS:")
        print(f"   🔧 Technical Skills ({len(skills_result.get('technical_skills', []))}): {skills_result.get('technical_skills', [])[:10]}{'...' if len(skills_result.get('technical_skills', [])) > 10 else ''}")
        print(f"   🤝 Soft Skills ({len(skills_result.get('soft_skills', []))}): {skills_result.get('soft_skills', [])[:10]}{'...' if len(skills_result.get('soft_skills', [])) > 10 else ''}")
        print(f"   🏢 Domain Skills ({len(skills_result.get('domain_skills', []))}): {skills_result.get('domain_skills', [])[:10]}{'...' if len(skills_result.get('domain_skills', [])) > 10 else ''}")
        print(f"   🎓 Certifications ({len(skills_result.get('certifications', []))}): {skills_result.get('certifications', [])}")
        print("="*50)
        
        return JSONResponse(
            status_code=200,
            content={
                "cv_filename": cv_filename,
                "skills": skills_result,
                "summary": {
                    "technical_count": len(skills_result.get("technical_skills", [])),
                    "soft_count": len(skills_result.get("soft_skills", [])),
                    "domain_count": len(skills_result.get("domain_skills", [])),
                    "certification_count": len(skills_result.get("certifications", []))
                }
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error processing CV: {str(e)}"}
        )

@app.post("/extract-skills-dynamic/")
async def extract_skills_dynamic(request: Request):
    """⚡ OPTIMIZED: Fast skill extraction using dynamic Claude-based extractor"""
    try:
        import time
        start_time = time.time()
        
        data = await request.json()
        mode = data.get("mode")  # "cv" or "jd"
        cv_filename = data.get("cv_filename", "")
        jd_url = data.get("jd_url", "")
        jd_text = data.get("jd_text", "")
        
        if mode not in ["cv", "jd"]:
            raise HTTPException(status_code=400, detail="Mode must be 'cv' or 'jd'")
        
        # ⚡ OPTIMIZATION 1: Reuse existing AI service instead of creating new client
        from .hybrid_ai_service import hybrid_ai
        import os
        
        # ⚡ OPTIMIZATION 2: Call functions directly instead of HTTP requests
        if mode == "cv":
            if not cv_filename:
                return {"error": "CV filename is required for CV mode"}
            
            # Get CV text directly from file system (no HTTP request)
            file_path = os.path.join(UPLOAD_DIR, cv_filename)
            if not os.path.exists(file_path):
                return {"error": "CV file not found"}
            
            # Extract text directly
            ext = os.path.splitext(cv_filename)[1].lower()
            if ext == ".pdf":
                from .cv_parser import extract_text_from_pdf
                input_text = extract_text_from_pdf(file_path)
            elif ext == ".docx":
                from .cv_parser import extract_text_from_docx
                input_text = extract_text_from_docx(file_path)
            else:
                return {"error": "Unsupported CV format"}
            
            text_type = "CV"
            
        elif mode == "jd":
            if jd_text:
                input_text = jd_text
            elif jd_url:
                # Use direct job scraper function (no HTTP request)
                from .job_scraper import scrape_job_description
                try:
                    input_text = scrape_job_description(jd_url)
                    if not input_text:
                        return {"error": "Failed to scrape job description"}
                except Exception as e:
                    return {"error": f"Failed to scrape JD: {str(e)}"}
            else:
                return {"error": "Either jd_text or jd_url is required for JD mode"}
            
            text_type = "Job Description"
        
        if not input_text or len(input_text.strip()) < 10:
            return {"error": "No valid text found to analyze"}
        
        # Use the enhanced Colab prompt and Claude Sonnet 4 for richer extraction
        prompt = f'''
Extract SOFT SKILLS, TECHNICAL SKILLS, and DOMAIN KEYWORDS from this {text_type.lower()} and categorize them:

IMPORTANT: Only extract skills/keywords that are explicitly mentioned or have very strong textual evidence. Avoid assumptions or industry-standard inferences. Do not repeat skills/keywords that are already mentioned.

## {text_type.upper()}:
{input_text}

## SOFT SKILLS:
EXPLICIT (directly stated): List soft skills clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List soft skills heavily suggested with strong textual evidence

## TECHNICAL SKILLS:
EXPLICIT (directly stated): List technical skills, tools, software, qualifications clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List technical skills heavily suggested with strong textual evidence

## DOMAIN KEYWORDS:
EXPLICIT:  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **industry terms**, **role-specific language**, **tools**, **methodologies**, **compliance concepts**, and **domain knowledge** that are **directly mentioned** in the {text_type.lower()}.

STRONGLY IMPLIED:
  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **domain-relevant concepts** that are **heavily suggested** by the context, responsibilities, or required outputs — even if not explicitly named.

### DO NOT include:
- Company or organization names
- Program/service titles
- Sector-level social causes
- Values or mission statements

## CONTEXT EVIDENCE:
For each skill/keyword, provide the relevant quote from the {text_type.lower()} that supports the extraction

At the end, provide THREE FINAL CLEAN LISTS:
1. SOFT SKILLS - Python list of strings containing all soft skills
2. TECHNICAL SKILLS - Python list of strings containing all technical skills
3. DOMAIN KEYWORDS - Python list of strings containing all domain-specific terms

Text: """
{input_text.strip()}
"""
'''
        
        # Use the enhanced Colab prompt and Claude Sonnet 4 for richer extraction
        prompt = f'''
Extract SOFT SKILLS, TECHNICAL SKILLS, and DOMAIN KEYWORDS from this job description and categorize them:

IMPORTANT: Only extract skills/keywords that are explicitly mentioned or have very strong textual evidence. Avoid assumptions or industry-standard inferences. Do not repeat skills/keywords that are already mentioned.

## JOB DESCRIPTION:
{input_text}

## SOFT SKILLS:
EXPLICIT (directly stated): List soft skills clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List soft skills heavily suggested with strong textual evidence

## TECHNICAL SKILLS:
EXPLICIT (directly stated): List technical skills, tools, software, qualifications clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List technical skills heavily suggested with strong textual evidence

## DOMAIN KEYWORDS:
EXPLICIT:  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **industry terms**, **role-specific language**, **tools**, **methodologies**, **compliance concepts**, and **domain knowledge** that are **directly mentioned** in the job description.

STRONGLY IMPLIED:
  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **domain-relevant concepts** that are **heavily suggested** by the context, responsibilities, or required outputs — even if not explicitly named.

### DO NOT include:
- Company or organization names
- Program/service titles
- Sector-level social causes
- Values or mission statements

## CONTEXT EVIDENCE:
For each skill/keyword, provide the relevant quote from the job description that supports the extraction

At the end, provide THREE FINAL CLEAN LISTS:
1. SOFT SKILLS - Python list of strings containing all soft skills
2. TECHNICAL SKILLS - Python list of strings containing all technical skills
3. DOMAIN KEYWORDS - Python list of strings containing all domain-specific terms

Text: """
{input_text.strip()}
"""
'''

        try:
            # Use DeepSeek for skill extraction
            from .hybrid_ai_service import hybrid_ai
            result_text = await hybrid_ai.generate_response(
                prompt=f"You are a precise extractor of skills from professional documents.\n\n{prompt}",
                max_tokens=1024,
                temperature=0.2
            )
            
            # Extract the three final clean lists from the response
            import re
            import ast
            result_text = result_text.strip()

            # The full formatted output for display
            comprehensive_analysis = result_text

            # Find all python lists in the response (should be 3)
            lists = re.findall(r'\[.*?\]', result_text, re.DOTALL)
            soft_skills, technical_skills, domain_keywords = [], [], []
            if len(lists) >= 3:
                try:
                    soft_skills = ast.literal_eval(lists[0])
                except Exception:
                    soft_skills = []
                try:
                    technical_skills = ast.literal_eval(lists[1])
                except Exception:
                    technical_skills = []
                try:
                    domain_keywords = ast.literal_eval(lists[2])
                except Exception:
                    domain_keywords = []
            else:
                # Fallback: try to parse by section headers
                soft_match = re.search(r'SOFT SKILLS.*?\n```python\n(.*?)```', result_text, re.DOTALL)
                tech_match = re.search(r'TECHNICAL SKILLS.*?\n```python\n(.*?)```', result_text, re.DOTALL)
                domain_match = re.search(r'DOMAIN KEYWORDS.*?\n```python\n(.*?)```', result_text, re.DOTALL)
                if soft_match:
                    try:
                        soft_skills = ast.literal_eval(soft_match.group(1))
                    except Exception:
                        soft_skills = []
                if tech_match:
                    try:
                        technical_skills = ast.literal_eval(tech_match.group(1))
                    except Exception:
                        technical_skills = []
                if domain_match:
                    try:
                        domain_keywords = ast.literal_eval(domain_match.group(1))
                    except Exception:
                        domain_keywords = []

            # If still empty, fallback to extracting from Markdown bullet points
            def extract_bullets(section_title):
                # Match section by title, then collect all lines starting with - until next ## or end
                pattern = rf'{section_title}:(.*?)(?:\n## |$)'
                match = re.search(pattern, result_text, re.DOTALL | re.IGNORECASE)
                if not match:
                    return []
                section = match.group(1)
                # Find all lines starting with -
                bullets = re.findall(r'^\s*[-\u2022]\s*(.+)', section, re.MULTILINE)
                # Remove quotes and explanations
                cleaned = []
                for b in bullets:
                    # Remove everything after ' - ' (keep only the skill name)
                    skill = b.split(' - ')[0].strip()
                    # Remove quotes and asterisks
                    skill = skill.strip('"'*2).strip("'* ")
                    if skill and skill.lower() not in ['n/a', 'none']:
                        cleaned.append(skill)
                return cleaned

            if not soft_skills:
                soft_skills = extract_bullets('SOFT SKILLS')
            if not technical_skills:
                technical_skills = extract_bullets('TECHNICAL SKILLS')
            if not domain_keywords:
                domain_keywords = extract_bullets('DOMAIN KEYWORDS')

            # Clean skill lists to remove empty, whitespace, and 'N/A' values
            def clean_skill_list(lst):
                return [s for s in lst if isinstance(s, str) and s.strip() and s.strip().lower() != 'n/a']
            soft_skills = clean_skill_list(soft_skills)
            technical_skills = clean_skill_list(technical_skills)
            domain_keywords = clean_skill_list(domain_keywords)

            end_time = time.time()
            
            # Log the comprehensive analysis to file
            
            # Print the full Claude AI analysis
            print("\n" + "="*80)
            print(f"📋 [{mode.upper()} CLAUDE ANALYSIS] Full Analysis Result:")
            print("="*80)
            print(comprehensive_analysis)
            print("="*80 + "\n")
            if mode == "cv":
                append_output_log(comprehensive_analysis, company_name="Maheshwor_Tiwari", tag="CV_CLAUDE_ANALYSIS")
            elif mode == "jd":
                append_output_log(comprehensive_analysis, company_name="Maheshwor_Tiwari", tag="JD_CLAUDE_ANALYSIS")
            print(f"⚡ [SONNET] Extraction completed in {end_time - start_time:.2f} seconds")
            return {
                "soft_skills": soft_skills,
                "technical_skills": technical_skills,
                "domain_keywords": domain_keywords,
                "comprehensive_analysis": comprehensive_analysis,
                "raw_response": result_text
            }
        except Exception as e:
            return {"error": f"AI extraction error: {str(e)}"}
            
    except Exception as e:
        return {"error": f"Error running skill extraction: {str(e)}"}

@app.post("/api/compare-skills")
async def compare_skills_endpoint(payload: dict = Body(...)):
    """
    Implements LLM-powered comparison between cached CV and JD skills.
    """
    import os
    import openai
    import json
    from fastapi import HTTPException

    # Prepare the prompt as per the user's specification
    prompt = f"""
You are performing a comparison analysis between pre-extracted and cached skill/keyword datasets. Your task is to categorize them into matched and missing categories with maximum precision.

## Input Data Structure
You will receive cached results from different sections where each contains data from both CV and job requirements:
- Cached Technical Skills: {payload.get('technical_skills', {})}
- Cached Soft Skills: {payload.get('soft_skills', {})}
- Cached Domain Keywords: {payload.get('domain_keywords', {})}

## Instructions
{'''
LLM-Powered Intelligent Matching
As an advanced language model, leverage your understanding to identify:
- Semantic Relationships (e.g., React.js, ReactJS, React → same technology)
- Skill Hierarchies (e.g., "Machine Learning" encompasses "Neural Networks")
- Industry Synonyms (e.g., "Customer Success" = "Client Relations")
- Contextual Equivalents (e.g., "API Development" matches "RESTful Services")
- Transferable Skills, Implied Competencies, Progressive Skills, Related Technologies
- Intelligent Inference: Understand context and nuance beyond exact string matching

For each job requirement, use your language understanding to determine:
- Match Confidence Levels (Exact, High, Medium, Low, No Match)
- Critical Thinking Process (Context Analysis, Skill Relationship Assessment, Practical Equivalence, Gap Impact Evaluation)

### Output Format
Return results in this exact JSON structure:
{
  "matched_technical_skills": [
    { "skill": "...", "match_type": "exact|semantic", "cv_reference": "..." }
  ],
  "matched_soft_skills": [
    { "skill": "...", "match_type": "exact|semantic", "cv_reference": "..." }
  ],
  "matched_domain_keywords": [
    { "keyword": "...", "match_type": "exact|semantic", "cv_reference": "..." }
  ],
  "missing_technical_skills": [
    { "skill": "..." }
  ],
  "missing_soft_skills": [
    { "skill": "..." }
  ],
  "missing_domain_keywords": [
    { "keyword": "..." }
  ]
}

### Precision Rules
- Avoid False Positives (e.g., Don't match "Java" with "JavaScript")
- Handle Ambiguity: If uncertain about semantic match, classify as missing
- Prefer exact matches over semantic matches
- When multiple CV skills could match one job requirement, choose the closest match
- Validation Checks: Ensure no skill appears in both matched and missing categories
- Confirm match_type accuracy (exact vs semantic)
- Special Cases: Technology Stacks, Experience Levels, Version Specificity
'''}

## Data
Technical Skills: {json.dumps(payload.get('technical_skills', {}))}
Soft Skills: {json.dumps(payload.get('soft_skills', {}))}
Domain Keywords: {json.dumps(payload.get('domain_keywords', {}))}

Return ONLY the JSON object in the required format.
"""

    # Call DeepSeek LLM
    try:
        from .hybrid_ai_service import hybrid_ai
        result_text = await hybrid_ai.generate_response(
            prompt=f"You are an expert ATS skill matcher.\n\n{prompt}",
            max_tokens=3600,
            temperature=0.2
        )

        # Extract JSON from the LLM response
        import re
        match = re.search(r'\{[\s\S]*\}', result_text)
        if match:
            json_str = match.group(0)
        else:
            json_str = result_text
        try:
            result = json.loads(json_str)
            
            # Remove duplicates from matched skills FIRST
            result = _consolidate_matched_skills(result)
            
            # CRITICAL VALIDATION: Check if all JD requirements are processed
            validation_result = _validate_comparison_completeness(result, jd_skills, total_jd_requirements)
            if not validation_result['valid']:
                print(f"⚠️ [VALIDATION] Incomplete processing detected: {validation_result['message']}")
                # Attempt to fix missing requirements
                result = _fix_missing_requirements(result, jd_skills, cv_skills, validation_result.get('missing_requirements', []))
            
            # Final validation and metrics
            final_validation = _validate_comparison_completeness(result, jd_skills, total_jd_requirements)
            
            # Ensure match_summary exists
            if 'match_summary' not in result:
                result['match_summary'] = {}
            
            result['match_summary']['validation_passed'] = final_validation['valid']
            result['match_summary']['processed_requirements'] = final_validation['processed_count']
            result['match_summary']['total_jd_requirements'] = total_jd_requirements
            
            print(f"✅ [VALIDATION] Final result: {final_validation['processed_count']}/{total_jd_requirements} requirements processed")
            
            return {
                "comparison_result": result,
                "raw_response": result_text,
                "validation": final_validation
            }
        except json.JSONDecodeError as e:
            return {
                "error": f"Failed to parse JSON: {e}",
                "raw_response": result_text
            }
    except Exception as e:
        return {"error": f"LLM comparison failed: {e}"}

@app.post("/api/llm/compare-skills")
async def compare_skills_llm(request: Request):
    """
    Enhanced intelligent skill comparison using AI matcher with detailed reasoning
    Inspired by Python implementation for superior semantic matching
    """
    try:
        data = await request.json()
        cv_skills = data.get("cv_skills", {})
        jd_skills = data.get("jd_skills", {})
        jd_text = data.get("jd_text", "")
        
        print(f"🚀 [API] Enhanced skill comparison endpoint called")
        print(f"📋 [API] CV Skills: {cv_skills}")
        print(f"📋 [API] JD Skills: {jd_skills}")
        print(f"📄 [API] JD Text length: {len(jd_text)} chars")
        
        # Extract company name from JD text for consistent file naming
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver()
        company_name = saver.extract_company_name(jd_text) if jd_text else "Company"
        
        print(f"🏢 [API] Extracted company name: {company_name}")
        
        # Use the enhanced AI matcher with reasoning
        from .ai_matcher import intelligent_skill_comparison
        
        # Get intelligent comparison with detailed reasoning - pass company name
        result = intelligent_skill_comparison(cv_skills, jd_skills, company_name=company_name)
        
        print(f"✅ [API] Enhanced comparison completed")
        print(f"📊 [API] Match summary: {result.get('match_summary', {})}")
        
        # Convert to the format expected by frontend
        formatted_result = _convert_to_frontend_format(result)
        
        return {
            "comparison_result": formatted_result,
            "enhanced_reasoning": True,
            "ai_powered": True
        }
        
    except Exception as e:
        print(f"❌ [API] Error in enhanced skill comparison: {e}")
        # Fallback to original implementation if enhanced fails
        return await _fallback_skill_comparison(cv_skills, jd_skills)

@app.post("/api/llm/generate-recommendations")
async def generate_ai_recommendations(request: Request):
    """
    Generate AI-powered CV tailoring recommendations based on analysis data
    """
    try:
        data = await request.json()
        prompt = data.get("prompt", "")
        cv_filename = data.get("cv_filename", "")
        jd_text = data.get("jd_text", "")
        analysis_text = data.get("analysis_text", "")
        
        print(f"🚀 [API] AI Recommendations generation endpoint called")
        print(f"📋 [API] CV: {cv_filename}")
        print(f"📋 [API] JD length: {len(jd_text)} characters")
        print(f"🧪 [API] Analysis text provided: {bool(analysis_text)}")
        
        if not cv_filename or not jd_text:
            raise HTTPException(status_code=400, detail="CV filename and JD text are required")
        
        session_filepath = None
        # Try session file manager (optional)
        try:
            from .session_file_manager import session_file_manager
            print(f"✅ [API] Session file manager imported successfully")
            try:
                session_filepath = session_file_manager.find_existing_session_file(cv_filename, jd_text)
                print(f"🔍 [API] Session file search completed: {bool(session_filepath)}")
            except Exception as search_error:
                print(f"⚠️ [API] Session file search error (non-fatal): {search_error}")
        except Exception as import_error:
            print(f"⚠️ [API] Session file manager not available (non-fatal): {import_error}")
        
        if session_filepath:
            print(f"📁 [API] Found session file: {session_filepath}")
            # Read the session file content
            session_data = session_file_manager.get_session_file_content(session_filepath)
            if session_data:
                # Extract analysis data from the session file
                extracted_text = _extract_analysis_from_session(session_data)
                if extracted_text:
                    analysis_text = extracted_text
                    print(f"📊 [API] Using analysis extracted from session file ({len(analysis_text)} chars)")
        
        # Print the input file name used for recommendations
        try:
            if session_filepath:
                print(f"📄 [API] Recommendations input file: {os.path.basename(session_filepath)}")
            else:
                print(f"📄 [API] Recommendations input file: {cv_filename} (no session file; using inline analysis)")
        except Exception as e:
            print(f"⚠️ [API] Could not print recommendations input file: {e}")
        
        if not analysis_text:
            raise HTTPException(status_code=400, detail="Analysis text is required (provide analysis_text or run prior steps)")
        
        # Use the hybrid AI service to generate recommendations
        from .hybrid_ai_service import hybrid_ai
        
        # Combine prompt and analysis text
        full_prompt = f"{prompt}\n\n## ANALYSIS DATA\n{analysis_text}"
        
        print(f"🤖 [API] Generating AI recommendations with Claude...")
        
        # Generate recommendations using Claude
        recommendations = hybrid_ai.generate_response(
            prompt=full_prompt,
            max_tokens=4000,
            temperature=0.3,
            provider="claude"
        )
        
        if recommendations:
            # Update session file with AI recommendations
            # Update session file if available
            if session_filepath:
                try:
                    session_file_manager.update_session_file(
                        session_filepath, 
                        "ai_recommendations", 
                        {
                            "recommendations": recommendations,
                            "generated_at": datetime.now().isoformat(),
                            "prompt_used": prompt[:200] + "..." if len(prompt) > 200 else prompt
                        }
                    )
                except Exception as sess_err:
                    print(f"⚠️ [API] Could not update session file: {sess_err}")
            else:
                print("ℹ️ [API] No session file found; skipping session update")
            
            # Log the recommendations
            append_output_log(
                recommendations, 
                company_name="AI_Recommendations", 
                tag="CV_TAILORING_RECOMMENDATIONS"
            )
            
            print(f"✅ [API] AI Recommendations generated successfully")
            print(f"📊 [API] Recommendations length: {len(recommendations)} characters")
            
            response = {
                "recommendations": recommendations,
                "success": True,
                "ai_powered": True
            }
            if session_filepath:
                try:
                    response["session_file"] = os.path.basename(session_filepath)
                except Exception:
                    pass
            return response
        else:
            raise HTTPException(status_code=500, detail="AI generation failed: No recommendations generated")
        
    except Exception as e:
        print(f"❌ [API] Error in AI recommendations generation: {e}")
        raise HTTPException(status_code=500, detail=f"AI recommendations generation failed: {e}")

def _extract_analysis_from_session(session_data: dict) -> str:
    """
    Extract comprehensive analysis text from session data for AI recommendations
    """
    buffer = []
    
    # Add session metadata
    metadata = session_data.get("session_metadata", {})
    buffer.append(f"## SESSION METADATA")
    buffer.append(f"CV Filename: {metadata.get('cv_filename', 'N/A')}")
    buffer.append(f"JD Preview: {metadata.get('jd_text_preview', 'N/A')}")
    buffer.append(f"Created: {metadata.get('created_at', 'N/A')}")
    buffer.append(f"Last Updated: {metadata.get('last_updated', 'N/A')}")
    buffer.append(f"Analysis Phase: {metadata.get('analysis_phase', 'N/A')}")
    buffer.append("")
    
    # Extract Preliminary Analysis
    prelim_data = session_data.get("analysis_results", {}).get("preliminary_analysis", {})
    if prelim_data:
        buffer.append("## PRELIMINARY ANALYSIS RESULTS")
        if "cv_skills" in prelim_data:
            cv_skills = prelim_data["cv_skills"]
            buffer.append("### CV Skills Extracted:")
            buffer.append(f"Technical Skills: {', '.join(cv_skills.get('technical_skills', []))}")
            buffer.append(f"Soft Skills: {', '.join(cv_skills.get('soft_skills', []))}")
            buffer.append(f"Domain Keywords: {', '.join(cv_skills.get('domain_keywords', []))}")
            buffer.append("")
        
        if "jd_skills" in prelim_data:
            jd_skills = prelim_data["jd_skills"]
            buffer.append("### Job Description Requirements:")
            buffer.append(f"Required Technical Skills: {', '.join(jd_skills.get('technical_skills', []))}")
            buffer.append(f"Required Soft Skills: {', '.join(jd_skills.get('soft_skills', []))}")
            buffer.append(f"Required Domain Keywords: {', '.join(jd_skills.get('domain_keywords', []))}")
            buffer.append("")
    
    # Extract Skill Comparison
    skill_comp_data = session_data.get("analysis_results", {}).get("skill_comparison", {})
    if skill_comp_data:
        buffer.append("## SKILL COMPARISON RESULTS")
        for skill_type in ["technical_skills", "soft_skills", "domain_keywords"]:
            if skill_type in skill_comp_data:
                skill_data = skill_comp_data[skill_type]
                buffer.append(f"### {skill_type.replace('_', ' ').title()}:")
                buffer.append(f"Matched: {', '.join(skill_data.get('matched', []))}")
                buffer.append(f"Missing: {', '.join(skill_data.get('missing', []))}")
                buffer.append("")
    
    # Extract Enhanced ATS Score
    ats_data = session_data.get("analysis_results", {}).get("enhanced_ats_score", {})
    if ats_data:
        buffer.append("## ENHANCED ATS SCORE ANALYSIS")
        buffer.append(f"Overall ATS Score: {ats_data.get('overall_ats_score', 'N/A')}/100")
        
        # Add detailed breakdown if available
        detailed_breakdown = ats_data.get("detailed_score_breakdown", {})
        if detailed_breakdown:
            buffer.append("### Detailed Score Breakdown:")
            for category, score in detailed_breakdown.items():
                buffer.append(f"{category}: {score}")
            buffer.append("")
        
        # Add skills analysis if available
        skills_analysis = ats_data.get("skills_analysis", {})
        if skills_analysis:
            buffer.append("### Skills Analysis:")
            cv_skills = skills_analysis.get("cv_skills", {})
            jd_skills = skills_analysis.get("jd_skills", {})
            
            if cv_skills:
                buffer.append("CV Skills:")
                buffer.append(f"Technical: {', '.join(cv_skills.get('technical_skills', []))}")
                buffer.append(f"Soft: {', '.join(cv_skills.get('soft_skills', []))}")
                buffer.append(f"Domain: {', '.join(cv_skills.get('domain_keywords', []))}")
                buffer.append("")
            
            if jd_skills:
                buffer.append("JD Requirements:")
                buffer.append(f"Technical: {', '.join(jd_skills.get('technical_skills', []))}")
                buffer.append(f"Soft: {', '.join(jd_skills.get('soft_skills', []))}")
                buffer.append(f"Domain: {', '.join(jd_skills.get('domain_keywords', []))}")
                buffer.append("")
        
        # Add skills matching if available
        skills_matching = ats_data.get("skills_matching", {})
        if skills_matching:
            buffer.append("### Skills Matching:")
            matched_skills = skills_matching.get("matched_skills", {})
            missing_skills = skills_matching.get("missing_skills", {})
            
            if matched_skills:
                buffer.append("Matched Skills:")
                for skill_type, skills in matched_skills.items():
                    if skills:
                        buffer.append(f"{skill_type}: {', '.join(skills)}")
                buffer.append("")
            
            if missing_skills:
                buffer.append("Missing Skills:")
                for skill_type, skills in missing_skills.items():
                    if skills:
                        buffer.append(f"{skill_type}: {', '.join(skills)}")
                buffer.append("")
        
        # Add market reality check if available
        market_reality = ats_data.get("market_reality_check", {})
        if market_reality:
            buffer.append("### Market Reality Check:")
            buffer.append(f"Market Flexibility: {market_reality.get('market_flexibility', 'N/A')}")
            buffer.append(f"Hiring Urgency: {market_reality.get('hiring_urgency', 'N/A')}")
            buffer.append(f"Competition Level: {market_reality.get('competition_level', 'N/A')}")
            buffer.append("")
        
        # Add strategic positioning if available
        strategic_positioning = ats_data.get("strategic_positioning", {})
        if strategic_positioning:
            buffer.append("### Strategic Positioning:")
            buffer.append(f"Primary Strategy: {strategic_positioning.get('primary_strategy', 'N/A')}")
            buffer.append(f"Key Advantages: {', '.join(strategic_positioning.get('key_advantages', []))}")
            buffer.append(f"Positioning Notes: {strategic_positioning.get('positioning_notes', 'N/A')}")
            buffer.append("")
    
    return "\n".join(buffer)

def _convert_to_frontend_format(ai_result: dict) -> dict:
    """
    Convert the enhanced AI matcher result to frontend-expected format
    Maps from Python-style result to Flutter app format
    """
    try:
        # Frontend expects this structure with matched/missing per category
        formatted = {
            "matched": {},
            "missing": {},
            "match_summary": ai_result.get('match_summary', {})
        }
        
        # Map each category from AI result to frontend format
        for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
            if category in ai_result:
                category_data = ai_result[category]
                
                # Convert matched skills with reasoning
                formatted["matched"][category] = []
                for match in category_data.get('matched', []):
                    formatted["matched"][category].append({
                        "cv_skill": match.get('cv_equivalent', ''),
                        "jd_requirement": match.get('jd_skill', ''),
                        "match_reason": match.get('reasoning', 'AI semantic match'),
                        "match_type": "semantic",  # All enhanced matches are semantic
                        "reasoning": match.get('reasoning', '')  # Extra reasoning field
                    })
                
                # Convert missing skills with reasoning
                formatted["missing"][category] = []
                for missing in category_data.get('missing', []):
                    formatted["missing"][category].append({
                        "jd_skill": missing.get('jd_skill', ''),
                        "reasoning": missing.get('reasoning', 'Not found in CV')
                    })
        
        print(f"🔄 [FORMAT] Converted AI result to frontend format")
        print(f"📊 [FORMAT] Categories: {list(formatted['matched'].keys())}")
        
        return formatted
        
    except Exception as e:
        print(f"❌ [FORMAT] Error converting format: {e}")
        # Return basic structure if conversion fails
        return {
            "matched": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
            "missing": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
            "match_summary": {"match_percentage": 0, "total_matched": 0, "total_missing": 0}
        }

async def _fallback_skill_comparison(cv_skills: dict, jd_skills: dict) -> dict:
    """
    Fallback to original skill comparison if enhanced AI matcher fails
    """
    try:
        print(f"🔄 [FALLBACK] Using fallback skill comparison")
        
        # Calculate exact requirement counts
        jd_tech_count = len(jd_skills.get('technical_skills', []))
        jd_soft_count = len(jd_skills.get('soft_skills', []))
        jd_domain_count = len(jd_skills.get('domain_keywords', []))
        total_jd_requirements = jd_tech_count + jd_soft_count + jd_domain_count
        
                # Basic prompt for fallback
        prompt = f"""
Compare CV skills against JD requirements and return JSON.

CV SKILLS:
Technical: {', '.join(cv_skills.get('technical_skills', []))}
Soft: {', '.join(cv_skills.get('soft_skills', []))}
Domain: {', '.join(cv_skills.get('domain_keywords', []))}

JD REQUIREMENTS:
Technical: {', '.join(jd_skills.get('technical_skills', []))}
Soft: {', '.join(jd_skills.get('soft_skills', []))}
Domain: {', '.join(jd_skills.get('domain_keywords', []))}

Return JSON with matched and missing skills for each category.
"""

        # Use DeepSeek as fallback
        from .hybrid_ai_service import hybrid_ai
        
        result_text = await hybrid_ai.generate_response(
            prompt=f"You are an ATS skill matcher. Return valid JSON.\n\n{prompt}",
            max_tokens=2000,
            temperature=0.1
        )
        
        # Try to parse JSON
        import re
        import json
        
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
            result = json.loads(json_str)
            
            # Add basic match summary
            if 'match_summary' not in result:
                total_matched = 0
                total_missing = 0
                
                for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
                    if 'matched' in result and category in result['matched']:
                        total_matched += len(result['matched'][category])
                    if 'missing' in result and category in result['missing']:
                        total_missing += len(result['missing'][category])
                
                match_percentage = (total_matched / (total_matched + total_missing) * 100) if (total_matched + total_missing) > 0 else 0
                
                result['match_summary'] = {
                    'total_matched': total_matched,
                    'total_missing': total_missing,
                    'match_percentage': round(match_percentage, 1),
                    'fallback_used': True
                }
            
            return {
                "comparison_result": result,
                "fallback_used": True,
                "raw_response": result_text
            }
        else:
            # If JSON parsing fails, return empty structure
            return {
                "comparison_result": {
                    "matched": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
                    "missing": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
                    "match_summary": {"match_percentage": 0, "fallback_failed": True}
                },
                "error": "Fallback JSON parsing failed",
                "raw_response": result_text
            }
            
    except Exception as e:
        print(f"❌ [FALLBACK] Fallback comparison failed: {e}")
        return {
            "comparison_result": {
                "matched": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
                "missing": {"technical_skills": [], "soft_skills": [], "domain_keywords": []},
                "match_summary": {"match_percentage": 0, "fallback_error": str(e)}
            },
            "error": f"Fallback failed: {str(e)}"
        }

def _consolidate_matched_skills(result):
    """
    Remove duplicate CV skills from matched results and consolidate JD requirements.
    Each CV skill should appear only once, even if it matches multiple JD requirements.
    """
    try:
        matched = result.get('matched', {})
        consolidated_matched = {}
        
        for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
            category_matches = matched.get(category, [])
            
            # Group by CV skill to remove duplicates
            cv_skill_groups = {}
            for match in category_matches:
                cv_skill = match.get('cv_skill', '')
                jd_requirement = match.get('jd_requirement', '')
                match_reason = match.get('match_reason', '')
                
                if cv_skill not in cv_skill_groups:
                    cv_skill_groups[cv_skill] = {
                        'cv_skill': cv_skill,
                        'jd_requirements': [],
                        'match_reasons': []
                    }
                
                cv_skill_groups[cv_skill]['jd_requirements'].append(jd_requirement)
                cv_skill_groups[cv_skill]['match_reasons'].append(match_reason)
            
            # Create consolidated matches with unique CV skills
            consolidated_category = []
            for cv_skill, group in cv_skill_groups.items():
                # Use the first JD requirement and combine match reasons if multiple
                primary_jd = group['jd_requirements'][0] if group['jd_requirements'] else ''
                
                if len(group['jd_requirements']) > 1:
                    # Multiple JD requirements matched to this CV skill
                    combined_reason = f"Matches {len(group['jd_requirements'])} requirements: {', '.join(group['jd_requirements'])}"
                else:
                    combined_reason = group['match_reasons'][0] if group['match_reasons'] else ''
                
                consolidated_category.append({
                    'cv_skill': cv_skill,
                    'jd_requirement': primary_jd,
                    'match_reason': combined_reason
                })
            
            consolidated_matched[category] = consolidated_category
        
        # Update the result with consolidated matches
        result['matched'] = consolidated_matched
        
        # Recalculate match summary with unique CV skills
        total_unique_matches = sum(len(consolidated_matched[cat]) for cat in consolidated_matched)
        total_jd_requirements = result.get('match_summary', {}).get('total_jd_requirements', 0)
        
        if total_jd_requirements > 0:
            match_percentage = round((total_unique_matches / total_jd_requirements) * 100)
        else:
            match_percentage = 0
        
        result['match_summary']['total_matches'] = total_unique_matches
        result['match_summary']['match_percentage'] = match_percentage
        
        print(f"✨ [Consolidation] Removed duplicates: {total_unique_matches} unique CV skills match {total_jd_requirements} JD requirements")
        
        return result
        
    except Exception as e:
        print(f"❌ [Consolidation] Error consolidating skills: {e}")
        return result  # Return original result if consolidation fails

def _validate_comparison_completeness(result, jd_skills, total_jd_requirements):
    """
    Validates if all JD requirements are accounted for in the comparison result.
    Returns detailed information about what's missing.
    """
    # Collect all JD requirements that should be processed
    all_jd_requirements = []
    for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
        all_jd_requirements.extend(jd_skills.get(category, []))
    
    print(f"🔍 [VALIDATION] All JD requirements to check: {all_jd_requirements}")
    
    # Collect matched requirements
    matched_jd_requirements = set()
    matched_section = result.get('matched', {})
    for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
        for match in matched_section.get(category, []):
            if isinstance(match, dict):
                matched_jd_requirements.add(match.get('jd_requirement'))
            else:
                print(f"⚠️ [VALIDATION] Unexpected match format: {match}")
    
    # Collect missing requirements  
    missing_from_result = []
    missing_section = result.get('missing', {})
    for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
        missing_from_result.extend(missing_section.get(category, []))
    
    # Find requirements that are completely unaccounted for
    accounted_requirements = matched_jd_requirements.union(set(missing_from_result))
    missing_jd_requirements = [req for req in all_jd_requirements if req not in accounted_requirements]
    
    processed_count = len(matched_jd_requirements) + len(missing_from_result)
    
    print(f"🔍 [VALIDATION] Matched: {len(matched_jd_requirements)}, Missing in result: {len(missing_from_result)}, Unaccounted: {len(missing_jd_requirements)}")
    print(f"🔍 [VALIDATION] Unaccounted requirements: {missing_jd_requirements}")

    if missing_jd_requirements:
        return {
            'valid': False,
            'message': f"Missing JD requirements: {', '.join(missing_jd_requirements)}",
            'missing_count': len(missing_jd_requirements),
            'missing_requirements': missing_jd_requirements,
            'processed_count': processed_count
        }
    else:
        return {
            'valid': True,
            'message': "All JD requirements accounted for.",
            'processed_count': processed_count,
            'missing_requirements': []
        }

def _fix_missing_requirements(result, jd_skills, cv_skills, missing_jd_requirements):
    """
    Attempts to fix missing JD requirements by adding them to the appropriate 'missing' category
    based on their original category in the JD skills.
    """
    print(f"🔧 [FIX] Fixing {len(missing_jd_requirements)} missing requirements: {missing_jd_requirements}")
    
    fixed_result = result.copy()
    fixed_matched = fixed_result.get('matched', {})
    fixed_missing = fixed_result.get('missing', {})
    
    # Ensure missing categories exist
    for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
        if category not in fixed_missing:
            fixed_missing[category] = []
    
    # Create mapping of requirement to original category
    requirement_categories = {}
    for category in ['technical_skills', 'soft_skills', 'domain_keywords']:
        for req in jd_skills.get(category, []):
            requirement_categories[req] = category
    
    # Add missing requirements to their correct categories
    for missing_req in missing_jd_requirements:
        # Check if already in matched (shouldn't be, but safety check)
        is_matched = False
        for category_matches in fixed_matched.values():
            if isinstance(category_matches, list):
                for match in category_matches:
                    if isinstance(match, dict) and match.get('jd_requirement') == missing_req:
                        is_matched = True
                        break
        
        if not is_matched:
            # Find the original category and add to missing
            original_category = requirement_categories.get(missing_req)
            if original_category and missing_req not in fixed_missing[original_category]:
                fixed_missing[original_category].append(missing_req)
                print(f"✅ [FIX] Added '{missing_req}' to missing {original_category}")
            elif not original_category:
                print(f"⚠️ [FIX] Could not find original category for '{missing_req}', adding to technical_skills")
                if missing_req not in fixed_missing['technical_skills']:
                    fixed_missing['technical_skills'].append(missing_req)
    
    # Update the result
    fixed_result['missing'] = fixed_missing
    
    # Recalculate totals
    total_matched = sum(len(fixed_matched.get(cat, [])) for cat in ['technical_skills', 'soft_skills', 'domain_keywords'])
    total_missing = sum(len(fixed_missing.get(cat, [])) for cat in ['technical_skills', 'soft_skills', 'domain_keywords'])
    total_processed = total_matched + total_missing
    
    # Update match summary
    if 'match_summary' not in fixed_result:
        fixed_result['match_summary'] = {}
    
    fixed_result['match_summary']['total_matches'] = total_matched
    fixed_result['match_summary']['processed_requirements'] = total_processed
    
    if fixed_result['match_summary'].get('total_jd_requirements', 0) > 0:
        fixed_result['match_summary']['match_percentage'] = round(
            (total_matched / fixed_result['match_summary']['total_jd_requirements']) * 100
        )
    
    print(f"✅ [FIX] Fixed result: {total_matched} matched + {total_missing} missing = {total_processed} total")
    return fixed_result

@app.post("/api/ats/enhanced-score")
async def calculate_enhanced_ats_score(request: Request):
    """
    Calculate Enhanced ATS Score using hybrid approach:
    - Existing keyword extraction as base
    - AI analysis for context, criticality, and intelligence
    """
    try:
        data = await request.json()
        cv_text = data.get("cv_text", "")
        jd_text = data.get("jd_text", "")
        skill_comparison = data.get("skill_comparison", {})
        extracted_keywords = data.get("extracted_keywords", {})
        
        if not all([cv_text, jd_text, skill_comparison, extracted_keywords]):
            raise HTTPException(status_code=400, detail="CV text, JD text, skill comparison, and extracted keywords are required")
        
        print(f"🚀 [ENHANCED_ATS] Starting Enhanced ATS Scoring...")
        print(f"   CV text length: {len(cv_text)} chars")
        print(f"   JD text length: {len(jd_text)} chars")
        print(f"   Skill comparison categories: {list(skill_comparison.keys())}")
        
        # Initialize Enhanced ATS Scorer (DeepSeek only)
        from .ats_enhanced_scorer import EnhancedATSScorer
        ats_scorer = EnhancedATSScorer()
        
        # Calculate enhanced ATS score
        enhanced_result = ats_scorer.calculate_enhanced_ats_score(
            cv_text=cv_text,
            jd_text=jd_text,
            skill_comparison=skill_comparison,
            extracted_keywords=extracted_keywords
        )
        
        # === UI-Style Pretty Print: Enhanced ATS Score ===
        print("\n" + "="*90)
        print("🎯 ATS Score Analysis - Final Results")
        print("="*90)
        
        # Overall Score Summary
        print(f"\n🎯 Overall Performance")
        print(f"Final ATS Score: {enhanced_result['overall_ats_score']}/100")
        print(f"Category: {enhanced_result['score_category']}")
        print("Needs improvement. Address critical gaps first.")
        
        # Base Scores Breakdown
        print(f"\n📊 Score Breakdown by Category")
        print(f"{'Category':<25}{'Score':<10}{'Status':<15}{'Impact':<10}")
        print("-" * 60)
        
        # Get scores from detailed breakdown since base_scores might be empty
        breakdown = enhanced_result.get('detailed_breakdown', {})
        base_scores = {
            'technical_skills': breakdown.get('technical_skills_match', {}).get('score', 0),
            'soft_skills': breakdown.get('soft_skills_match', {}).get('score', 0),
            'domain_keywords': breakdown.get('domain_keywords_match', {}).get('score', 0)
        }
        
        for cat, label in zip(['technical_skills', 'soft_skills', 'domain_keywords'],
                              ['Technical Skills', 'Soft Skills', 'Domain Keywords']):
            score = base_scores.get(cat, 0)
            status = "✅ Strong" if score >= 70 else "⚠️ Moderate" if score >= 40 else "❌ Weak"
            impact = "High" if score >= 70 else "Medium" if score >= 40 else "Low"
            print(f"{label:<25}{score:<10.1f}{status:<15}{impact:<10}")
        
        # Detailed Component Analysis
        print(f"\n🧠 Detailed Component Analysis")
        print(f"{'Component':<30}{'Score':<8}{'Weight':<8}{'Contribution':<10}")
        print("-" * 60)
        
        # Map component names to display names
        component_display_names = {
            'technical_skills_match': 'Technical Skills Match',
            'soft_skills_match': 'Soft Skills Match', 
            'domain_keywords_match': 'Domain Keywords Match',
            'skills_relevance': 'Skills Relevance',
            'experience_alignment': 'Experience Alignment',
            'industry_fit': 'Industry Fit',
            'role_seniority': 'Role Seniority',
            'technical_depth': 'Technical Depth',
            'criticality_bonus': 'Criticality Bonus',
            'requirement_bonus': 'Requirement Bonus'
        }
        
        for key, comp in breakdown.items():
            score = comp.get('score', 0)
            weight = comp.get('weight', 0)
            contribution = comp.get('contribution', 0)
            
            # Use display name if available, otherwise format the key
            display_name = component_display_names.get(key, key.replace('_', ' ').title())
            print(f"{display_name:<30}{score:<8.1f}{weight:<8.1f}{contribution:<10.1f}")
        
        # Critical Requirements Status
        print(f"\n🎯 Critical Requirements Status")
        print("-" * 60)
        
        # Get requirement bonus data from detailed breakdown
        critical_matches = 0
        critical_total = 0
        for rec in enhanced_result.get('recommendations', []):
            if "4/6 essential requirements met" in rec.lower():
                critical_matches = 4
                critical_total = 6
                break
        
        print(f"Essential Requirements Met: {critical_matches}/{critical_total} ✅")
        print(f"Preferred Requirements Met: None")
        
        # Get missing skills from recommendations
        missing_skills = []
        for rec in enhanced_result.get('recommendations', []):
            if "URGENT: Add critical missing skills:" in rec:
                skills_text = rec.split("URGENT: Add critical missing skills:")[1].strip()
                missing_skills = [s.strip() for s in skills_text.split(",")]
        
        if missing_skills:
            print(f"Missing Critical Skills: {', '.join(missing_skills)}")
        
        # Get matched skills from base scores
        matched_skills = []
        if base_scores.get('technical_skills', 0) > 60:
            matched_skills.append("Business Intelligence tools")
        if base_scores.get('domain_keywords', 0) > 60:
            matched_skills.append("Databases")
        
        if matched_skills:
            print(f"Key Strength: {', '.join(matched_skills)}")
        
        print("="*90 + "\n")

        # Capture the output and append to log file
        from .print_output_logger import append_output_log
        
        output_text = f"""
🎯 ATS Score Analysis - Final Results
{'='*90}

🎯 Overall Performance
Final ATS Score: {enhanced_result['overall_ats_score']}/100
Category: {enhanced_result['score_category']}
Needs improvement. Address critical gaps first.

📊 Score Breakdown by Category
{'Category':<25}{'Score':<10}{'Status':<15}{'Impact':<10}
{'-' * 60}"""

        # Add base scores
        for cat, label in zip(['technical_skills', 'soft_skills', 'domain_keywords'],
                              ['Technical Skills', 'Soft Skills', 'Domain Keywords']):
            score = base_scores.get(cat, 0)
            status = "✅ Strong" if score >= 70 else "⚠️ Moderate" if score >= 40 else "❌ Weak"
            impact = "High" if score >= 70 else "Medium" if score >= 40 else "Low"
            output_text += f"\n{label:<25}{score:<10.1f}{status:<15}{impact:<10}"

        # Add component analysis
        output_text += f"""

🧠 Detailed Component Analysis
{'Component':<30}{'Score':<8}{'Weight':<8}{'Contribution':<10}
{'-' * 60}"""

        for key, comp in breakdown.items():
            score = comp.get('score', 0)
            weight = comp.get('weight', 0)
            contribution = comp.get('contribution', 0)
            display_name = component_display_names.get(key, key.replace('_', ' ').title())
            output_text += f"\n{display_name:<30}{score:<8.1f}{weight:<8.1f}{contribution:<10.1f}"

        # Add critical requirements
        output_text += f"""

🎯 Critical Requirements Status
{'-' * 60}
Essential Requirements Met: {critical_matches}/{critical_total} ✅
Preferred Requirements Met: None"""

        if missing_skills:
            output_text += f"\nMissing Critical Skills: {', '.join(missing_skills)}"
        
        if matched_skills:
            output_text += f"\nKey Strength: {', '.join(matched_skills)}"

        output_text += f"\n{'='*90}\n"

        # Append to the log file with company name extracted from JD
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver()
        company_name = saver.extract_company_name(jd_text)
        append_output_log(output_text, company_name=company_name, tag="ATS_SCORE")
        
        # The ATS results are now automatically appended to the main analysis file via append_output_log()
        # This ensures all analysis steps (preliminary analysis, skill comparison, ATS test) are in one file
        analysis_filename = f"{company_name}_output_log.txt"
        enhanced_result["analysis_file_saved"] = analysis_filename
        print(f"✅ [ENHANCED_ATS] ATS results appended to main analysis file: {analysis_filename}")
        
        return enhanced_result
        
    except Exception as e:
        print(f"❌ [ENHANCED_ATS] Error: {str(e)}")
        logger.error(f"Enhanced ATS scoring error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Enhanced ATS scoring failed: {str(e)}")


@app.post("/api/save-analysis-results")
async def save_analysis_results(request: Request):
    """
    Save comprehensive analysis results to a text file named based on JD company name
    """
    try:
        data = await request.json()
        cv_text = data.get("cv_text", "")
        jd_text = data.get("jd_text", "")
        skill_comparison = data.get("skill_comparison", {})
        ats_results = data.get("ats_results", {})
        company_name = data.get("company_name", None)
        
        print(f"💾 [SAVE_ANALYSIS] Starting analysis results save...")
        print(f"   CV text length: {len(cv_text)} chars")
        print(f"   JD text length: {len(jd_text)} chars")
        print(f"   Skill comparison keys: {list(skill_comparison.keys()) if skill_comparison else 'None'}")
        print(f"   ATS results keys: {list(ats_results.keys()) if ats_results else 'None'}")
        print(f"   Company name: {company_name}")
        
        # Check each required field individually for better error reporting
        missing_fields = []
        if not cv_text:
            missing_fields.append("cv_text")
        if not jd_text:
            missing_fields.append("jd_text")
        
        # Be more flexible about skill_comparison and ats_results
        # They can be empty objects or contain minimal data
        if skill_comparison is None:
            missing_fields.append("skill_comparison")
        if ats_results is None:
            missing_fields.append("ats_results")
        
        if missing_fields:
            error_msg = f"Missing required fields: {', '.join(missing_fields)}"
            print(f"❌ [SAVE_ANALYSIS] Validation error: {error_msg}")
            raise HTTPException(status_code=400, detail=error_msg)
        
        # Ensure we have at least some data for skill_comparison and ats_results
        if not skill_comparison:
            skill_comparison = {"status": "completed", "message": "No detailed comparison available"}
        if not ats_results:
            ats_results = {"status": "completed", "message": "No detailed ATS results available"}
        
        # Initialize Simplified Analysis Results Saver
        from .analysis_results_saver_simple import SimpleAnalysisResultsSaver
        saver = SimpleAnalysisResultsSaver(debug=True)
        
        # Save analysis results
        filepath = saver.save_analysis_results(
            cv_text=cv_text,
            jd_text=jd_text,
            skill_comparison=skill_comparison,
            ats_results=ats_results,
            company_name=company_name
        )
        
        # Get filename from path
        filename = os.path.basename(filepath)
        
        print(f"✅ [SAVE_ANALYSIS] Analysis results saved: {filename}")
        
        return {
            "message": "Analysis results saved successfully",
            "filename": filename,
            "filepath": filepath
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"❌ [SAVE_ANALYSIS] Error: {str(e)}")
        import traceback
        print(f"❌ [SAVE_ANALYSIS] Full traceback:")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to save analysis results: {str(e)}")


@app.post("/api/generate-recommendations")
async def generate_recommendations_from_file(request: Request):
    """
    Generate LLM-based recommendations from a saved analysis file
    """
    try:
        data = await request.json()
        analysis_filepath = data.get("analysis_filepath", "")
        
        if not analysis_filepath:
            raise HTTPException(status_code=400, detail="Analysis file path is required")
        
        # Check if file exists
        if not os.path.exists(analysis_filepath):
            raise HTTPException(status_code=404, detail="Analysis file not found")
        
        print(f"🧠 [RECOMMENDATIONS] Generating recommendations from: {analysis_filepath}")
        
        # Get API key
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")
        
        # Initialize Analysis Results Saver with debug enabled
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Generate recommendations
        recommendations = saver.generate_llm_recommendations(
            analysis_filepath=analysis_filepath,
            api_key=api_key
        )
        
        print(f"✅ [RECOMMENDATIONS] Generated recommendations successfully")
        
        return {
            "message": "Recommendations generated successfully",
            "recommendations": recommendations,
            "source_file": analysis_filepath
        }
        
    except Exception as e:
        print(f"❌ [RECOMMENDATIONS] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate recommendations: {str(e)}")


@app.post("/api/llm/generate-recommendations-from-analysis")
async def generate_recommendations_from_latest_analysis(request: Request):
    """
    Generate LLM-based recommendations from the latest saved analysis file automatically.
    This endpoint is designed to work with the ATS tab frontend.
    """
    try:
        data = await request.json()
        cv_filename = data.get("cv_filename", "")
        jd_text = data.get("jd_text", "")
        custom_prompt = data.get("prompt", "")  # Frontend can provide custom prompt
        
        print(f"🧠 [AUTO_RECOMMENDATIONS] Starting recommendations generation...")
        print(f"   CV: {cv_filename}")
        print(f"   JD text length: {len(jd_text)} chars")
        
        # Get the analysis file for the specific CV+JD combination
        analysis_dir = "analysis_results"
        if not os.path.exists(analysis_dir):
            raise HTTPException(status_code=404, detail="No analysis results directory found")
        
        # Extract company name from JD to find the correct analysis file
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver()
        company_name = saver.extract_company_name(jd_text)
        
        # Clean company name for folder structure
        company_name_clean = re.sub(r'[^\w\s&.-]', '', company_name)
        company_name_clean = re.sub(r'\s+', '_', company_name_clean)
        company_name_clean = company_name_clean.strip('_')
        
        # Look for the specific analysis log file for this JD
        specific_analysis_file = f"{company_name_clean}_output_log.txt"
        specific_analysis_path = os.path.join(analysis_dir, company_name_clean, specific_analysis_file)
        
        if os.path.exists(specific_analysis_path):
            # Use the specific analysis file for this CV+JD combination
            analysis_filepath = specific_analysis_path
            latest_file = specific_analysis_file
            print(f"📁 [AUTO_RECOMMENDATIONS] Using specific analysis file: {specific_analysis_file}")
        else:
            # Fallback: Get all analysis files from subdirectories and use the most recent
            analysis_files = []
            for company_dir in os.listdir(analysis_dir):
                company_path = os.path.join(analysis_dir, company_dir)
                if os.path.isdir(company_path):
                    for file in os.listdir(company_path):
                        if file.endswith("_output_log.txt"):
                            full_path = os.path.join(company_path, file)
                            analysis_files.append(full_path)
            
            if not analysis_files:
                raise HTTPException(status_code=404, detail="No analysis files found. Please run the analysis cycle first.")
            
            # Sort by modification time (most recent first)
            analysis_files.sort(key=lambda x: os.path.getmtime(x), reverse=True)
            analysis_filepath = analysis_files[0]
            latest_file = os.path.basename(analysis_filepath)
            print(f"📁 [AUTO_RECOMMENDATIONS] Using fallback analysis file: {latest_file}")
        
        # Use DeepSeek for LLM recommendations
        
        # Read the analysis file content
        with open(analysis_filepath, 'r', encoding='utf-8') as f:
            analysis_content = f.read()
        
        # Import the centralized prompt module
        try:
            from .ai_recommendations import get_ai_recommendations_prompt
        except ImportError:
            # Fallback if the prompts module is not available
            def get_ai_recommendations_prompt(analysis_content):
                return f"""
You are an expert CV tailoring consultant. Analyze the provided comprehensive analysis data and generate actionable CV tailoring recommendations.

COMPREHENSIVE ANALYSIS DATA:
{analysis_content}

Based on this detailed analysis, provide specific, actionable recommendations for improving the CV to better match the job requirements. Focus on:

1. **Technical Skills Enhancement**: Based on the skills comparison and gaps identified
2. **Soft Skills Positioning**: Leverage identified strengths and address weaknesses  
3. **ATS Score Optimization**: Address specific ATS scoring factors and requirements
4. **Strategic Positioning**: Based on the match analysis and market insights
5. **Keyword Integration**: Use the domain keywords analysis for strategic placement

Provide concrete, implementable advice that will maximize the candidate's chances of success.
"""

        # Use custom prompt if provided, otherwise use the centralized unified prompt
        if custom_prompt:
            # Frontend provided custom prompt - use it with the analysis content
            final_prompt = f"{custom_prompt}\n\nANALYSIS DATA:\n{analysis_content}"
        else:
            # Use centralized unified prompt
            final_prompt = get_ai_recommendations_prompt(analysis_content)
        
        # Generate recommendations using DeepSeek
        from .hybrid_ai_service import hybrid_ai
        
        try:
            recommendations = await hybrid_ai.generate_response(
                prompt=final_prompt,
                model="deepseek-chat",
                max_tokens=3000,
                temperature=0.1
            )
        except Exception as llm_error:
            print(f"❌ [AUTO_RECOMMENDATIONS] DeepSeek API error: {llm_error}")
            raise HTTPException(status_code=500, detail=f"DeepSeek API error: {llm_error}")
        
        print(f"✅ [AUTO_RECOMMENDATIONS] Generated recommendations successfully")
        print(f"   Recommendations length: {len(recommendations)} chars")
        
        # Append recommendations to analysis file
        try:
            append_result = append_output_log(recommendations, company_name=company_name, tag="AI_RECOMMENDATIONS")
            if append_result:
                print(f"📝 [AUTO_RECOMMENDATIONS] Results saved to analysis file")
            else:
                print(f"⚠️ [AUTO_RECOMMENDATIONS] Failed to save results to analysis file")
        except Exception as append_error:
            print(f"❌ [AUTO_RECOMMENDATIONS] Error saving to analysis file: {append_error}")
        
        return {
            "recommendations": recommendations,
            "source_file": latest_file,
            "analysis_filepath": analysis_filepath,
            "message": "Recommendations generated from latest analysis file"
        }
        
    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [AUTO_RECOMMENDATIONS] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate recommendations: {str(e)}")





if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
@app.get("/api/get-latest-analysis-file")
async def get_latest_analysis_file():
    """
    Get the path to the most recent analysis file.
    """
    try:
        analysis_dir = "analysis_results"
        if not os.path.exists(analysis_dir):
            raise HTTPException(status_code=404, detail="No analysis results directory found")
        
        # Get all analysis files and use the most recent
        analysis_files = [f for f in os.listdir(analysis_dir) if f.endswith(".txt")]
        if not analysis_files:
            raise HTTPException(status_code=404, detail="No analysis files found")
        
        # Sort by modification time (most recent first)
        analysis_files.sort(key=lambda x: os.path.getmtime(os.path.join(analysis_dir, x)), reverse=True)
        latest_file = analysis_files[0]
        filepath = os.path.join(analysis_dir, latest_file)
        
        print(f"📁 [ANALYSIS] Using latest analysis file: {filepath}")
        return {"filepath": filepath, "filename": latest_file}

    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [ANALYSIS] Error getting latest analysis file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get latest analysis file: {str(e)}")





@app.post("/api/save-exact-ui-output")
async def save_exact_ui_output(request: Request):
    """
    Save exact UI output text as displayed in Flutter app to JSON file
    This captures every letter, word, emoji, and formatting exactly as shown in UI
    """
    try:
        data = await request.json()
        
        # Extract required fields
        step_name = data.get("step_name", "")
        ui_output_text = data.get("ui_output_text", "")
        company_name = data.get("company_name", "Company")
        timestamp = data.get("timestamp", datetime.now().isoformat())
        
        if not ui_output_text:
            raise HTTPException(status_code=400, detail="ui_output_text is required")
        
        print(f"💾 [SAVE_UI_TEXT] Saving exact UI output for step: {step_name}")
        print(f"📏 [SAVE_UI_TEXT] Text length: {len(ui_output_text)} characters")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Use the new method to save exact UI text
        filepath = saver.save_exact_ui_text(
            step_name=step_name,
            ui_output_text=ui_output_text,
            company_name=company_name,
            timestamp=timestamp
        )
        
        filename = os.path.basename(filepath)
        
        print(f"✅ [SAVE_UI_TEXT] Exact UI output saved: {filename}")
        
        return {
            "message": "Exact UI output saved successfully",
            "filename": filename,
            "filepath": filepath,
            "step_name": step_name,
            "characters_saved": len(ui_output_text)
        }
        
    except Exception as e:
        print(f"❌ [SAVE_UI_TEXT] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save exact UI output: {str(e)}")
        
    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [INTELLIGENT_ANALYSIS] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get intelligent skill analysis: {str(e)}")


@app.post("/api/save-ui-outputs")
async def save_ui_outputs(request: Request):
    """
    Save exact UI outputs from each step in JSON format
    """
    try:
        data = await request.json()
        
        # Get the outputs from each step
        analyze_match_output = data.get("analyze_match_output")
        skill_comparison_output = data.get("skill_comparison_output")
        ats_test_output = data.get("ats_test_output")
        company_name = data.get("company_name")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Save UI outputs
        filepath = saver.save_ui_outputs_to_json(
            analyze_match_output=analyze_match_output,
            skill_comparison_output=skill_comparison_output,
            ats_test_output=ats_test_output,
            company_name=company_name
        )
        
        print(f"✅ [UI_OUTPUTS] UI outputs saved: {filepath}")
        
        return {
            "message": "UI outputs saved successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [UI_OUTPUTS] Error: {str(e)}")
        logger.error(f"UI outputs save error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save UI outputs: {str(e)}")



@app.post("/api/save-skill-comparison-output")
async def save_skill_comparison_output(request: Request):
    """
    Save the exact output from ATS tab skill comparison
    """
    try:
        data = await request.json()
        
        # Get the skill comparison output
        skill_comparison_output = data.get("skill_comparison_output", {})
        company_name = data.get("company_name")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Save skill comparison output
        filepath = saver.save_ui_outputs_to_json(
            skill_comparison_output=skill_comparison_output,
            company_name=company_name
        )
        
        print(f"✅ [SKILL_COMPARISON] Skill comparison output saved: {filepath}")
        
        return {
            "message": "Skill comparison output saved successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [SKILL_COMPARISON] Error: {str(e)}")
        logger.error(f"Skill comparison output save error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save skill comparison output: {str(e)}")

@app.post("/api/save-ats-test-output")
async def save_ats_test_output(request: Request):
    """
    Save the exact output from ATS tab test (detailed score breakdown + enhanced requirement bonus)
    """
    try:
        data = await request.json()
        
        # Get the ATS test output
        ats_test_output = data.get("ats_test_output", {})
        company_name = data.get("company_name")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Save ATS test output
        filepath = saver.save_ui_outputs_to_json(
            ats_test_output=ats_test_output,
            company_name=company_name
        )
        
        print(f"✅ [ATS_TEST] ATS test output saved: {filepath}")
        
        return {
            "message": "ATS test output saved successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [ATS_TEST] Error: {str(e)}")
        logger.error(f"ATS test output save error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save ATS test output: {str(e)}")




@app.post("/api/auto-save-skill-comparison")
async def auto_save_skill_comparison(request: Request):
    """
    Automatically save skill comparison output when displayed in Flutter UI
    """
    try:
        data = await request.json()
        
        # Get the skill comparison output
        skill_comparison_output = data.get("skill_comparison_output", {})
        company_name = data.get("company_name")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Auto-save skill comparison output
        filepath = saver.auto_save_skill_comparison(
            skill_comparison_output=skill_comparison_output,
            company_name=company_name
        )
        
        print(f"✅ [AUTO_SAVE] Skill comparison output auto-saved: {filepath}")
        
        return {
            "message": "Skill comparison output auto-saved successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [AUTO_SAVE] Error: {str(e)}")
        logger.error(f"Auto-save skill comparison error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to auto-save skill comparison output: {str(e)}")

@app.post("/api/auto-save-ats-test")
async def auto_save_ats_test(request: Request):
    """
    Automatically save ATS test output when displayed in Flutter UI
    """
    try:
        data = await request.json()
        
        # Get the ATS test output
        ats_test_output = data.get("ats_test_output", {})
        company_name = data.get("company_name")
        
        # Initialize Analysis Results Saver
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Auto-save ATS test output
        filepath = saver.auto_save_ats_test(
            ats_test_output=ats_test_output,
            company_name=company_name
        )
        
        print(f"✅ [AUTO_SAVE] ATS test output auto-saved: {filepath}")
        
        return {
            "message": "ATS test output auto-saved successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [AUTO_SAVE] Error: {str(e)}")
        logger.error(f"Auto-save ATS test error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to auto-save ATS test output: {str(e)}")


@app.post("/api/append-analyze-match")
async def append_analyze_match(request: Request):
    """
    Simply append analyze match output to file
    """
    try:
        data = await request.json()
        raw_analysis = data.get("raw_analysis", "")
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.append_analyze_match(raw_analysis, company_name)
        
        return {
            "message": "Analyze match output appended successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [APPEND] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to append analyze match: {str(e)}")

@app.post("/api/append-skill-comparison")
async def append_skill_comparison(request: Request):
    """
    Simply append skill comparison output to file
    """
    try:
        data = await request.json()
        skill_comparison_text = data.get("skill_comparison_text", "")
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.append_skill_comparison(skill_comparison_text, company_name)
        
        return {
            "message": "Skill comparison output appended successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [APPEND] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to append skill comparison: {str(e)}")

@app.post("/api/append-ats-test")
async def append_ats_test(request: Request):
    """
    Simply append ATS test output to file
    """
    try:
        data = await request.json()
        ats_test_text = data.get("ats_test_text", "")
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.append_ats_test(ats_test_text, company_name)
        
        return {
            "message": "ATS test output appended successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [APPEND] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to append ATS test: {str(e)}")


@app.post("/api/append-detailed-skill-analysis")
async def append_detailed_skill_analysis(request: Request):
    """
    Append detailed skill analysis with full breakdown to file
    """
    try:
        data = await request.json()
        detailed_analysis = data.get("detailed_analysis", {})
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.append_detailed_skill_analysis(detailed_analysis, company_name)
        
        return {
            "message": "Detailed skill analysis appended successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [APPEND] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to append detailed skill analysis: {str(e)}")




@app.post("/api/save-skill-comparison-json")
async def save_skill_comparison_json(request: Request):
    """
    Save skill comparison output to JSON file
    """
    try:
        data = await request.json()
        skill_comparison_output = data.get("skill_comparison_output", {})
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.auto_save_skill_comparison_json(skill_comparison_output, company_name)
        
        return {
            "message": "Skill comparison saved to JSON successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [JSON SAVE] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save skill comparison to JSON: {str(e)}")

@app.post("/api/save-ats-score-json")
async def save_ats_score_json(request: Request):
    """
    Save ATS score output to JSON file
    """
    try:
        data = await request.json()
        ats_score_output = data.get("ats_score_output", {})
        company_name = data.get("company_name", "Company")
        
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        filepath = saver.auto_save_ats_score_json(ats_score_output, company_name)
        
        return {
            "message": "ATS score saved to JSON successfully",
            "filepath": filepath
        }
        
    except Exception as e:
        print(f"❌ [JSON SAVE] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save ATS score to JSON: {str(e)}")

@app.get("/api/get-complete-analysis-json/{company_name}")
async def get_complete_analysis_json(company_name: str):
    """
    Get the complete analysis JSON file for a company
    """
    try:
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver(debug=True)
        
        # Create company subdirectory
        company_dir = os.path.join(saver.results_dir, company_name)
        os.makedirs(company_dir, exist_ok=True)
        
        filename = f"Complete_Analysis.json"
        filepath = os.path.join(company_dir, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data
        else:
            raise HTTPException(status_code=404, detail=f"No analysis file found for {company_name}")
            
    except Exception as e:
        print(f"❌ [GET JSON] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get analysis JSON: {str(e)}")


@app.post("/preliminary-analysis/")
async def preliminary_analysis(request: Request):
    """🎯 PRELIMINARY ANALYSIS: Run both CV and JD skill extraction side-by-side"""
    try:
        import time
        start_time = time.time()
        
        data = await request.json()
        cv_filename = data.get("cv_filename", "")
        jd_text = data.get("jd_text", "")
        
        if not cv_filename:
            return {"error": "CV filename is required"}
        if not jd_text:
            return {"error": "JD text is required"}
        
        print("\n" + "="*80)
        print("🎯 PRELIMINARY ANALYSIS: Starting CV and JD skill extraction")
        print("="*80)
        print(f"CV Filename: {cv_filename}")
        print(f"JD Text Length: {len(jd_text)} characters")
        print("="*80)
        
        # ⚡ OPTIMIZATION: Reuse existing AI service
        from .hybrid_ai_service import hybrid_ai
        
        # Get CV text
        file_path = os.path.join(UPLOAD_DIR, cv_filename)
        if not os.path.exists(file_path):
            return {"error": "CV file not found"}
        
        ext = os.path.splitext(cv_filename)[1].lower()
        if ext == ".pdf":
            from .cv_parser import extract_text_from_pdf
            cv_text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            from .cv_parser import extract_text_from_docx
            cv_text = extract_text_from_docx(file_path)
        else:
            return {"error": "Unsupported CV format"}
        
        if not cv_text or len(cv_text.strip()) < 10:
            return {"error": "No valid CV text found to analyze"}
        
        # Normalize both CV and JD text for consistent processing
        def normalize_document_text(text: str, doc_type: str) -> str:
            """Normalize text to ensure consistent processing for both CV and JD"""
            # Clean up the text
            normalized = text.strip()
            
            # Log original length
            print(f"📝 [{doc_type}] Original text length: {len(text)} chars")
            
            # Truncate if too long (same limit for both CV and JD)
            max_chars = 8000  # Same limit for both documents
            if len(normalized) > max_chars:
                normalized = normalized[:max_chars]
                print(f"✂️ [{doc_type}] Text truncated to {max_chars} chars")
            
            print(f"📝 [{doc_type}] Final text length: {len(normalized)} chars")
            return normalized
        
        # Normalize both documents using the same process
        cv_text = normalize_document_text(cv_text, "CV")
        jd_text = normalize_document_text(jd_text, "JD")
        
        # Create unified skill extraction function
        def create_skill_extraction_prompt(document_type: str, document_text: str) -> str:
            """Create standardized prompt for both CV and JD extraction"""
            return f'''
Extract SOFT SKILLS, TECHNICAL SKILLS, and DOMAIN KEYWORDS from this {document_type.lower()} and categorize them:

IMPORTANT: Only extract skills/keywords that are explicitly mentioned or have very strong textual evidence. Avoid assumptions or industry-standard inferences. Do not repeat skills/keywords that are already mentioned.

## {document_type.upper()}:
{document_text}

## SOFT SKILLS:
EXPLICIT (directly stated): List soft skills clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List soft skills heavily suggested with strong textual evidence

## TECHNICAL SKILLS:
EXPLICIT (directly stated): List technical skills, tools, software, qualifications clearly mentioned
STRONGLY IMPLIED (very likely based on responsibilities): List technical skills heavily suggested with strong textual evidence

## DOMAIN KEYWORDS:
EXPLICIT:  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **industry terms**, **role-specific language**, **tools**, **methodologies**, **compliance concepts**, and **domain knowledge** that are **directly mentioned** in the {document_type.lower()}.

STRONGLY IMPLIED:
  Your task is to extract **domain-specific keywords** related strictly to the **job role or functional expertise**, NOT the industry, organization, or its values.
  List **domain-relevant concepts** that are **heavily suggested** by the context, responsibilities, or required outputs — even if not explicitly named.

### DO NOT include:
- Company or organization names
- Program/service titles
- Sector-level social causes
- Values or mission statements

## CONTEXT EVIDENCE:
For each skill/keyword, provide the relevant quote from the {document_type.lower()} that supports the extraction

**CRITICAL OUTPUT REQUIREMENT:**
You MUST end your response with EXACTLY these three Python lists (no extra text after them):

SOFT_SKILLS = ["skill1", "skill2", "skill3"]
TECHNICAL_SKILLS = ["skill1", "skill2", "skill3"]
DOMAIN_KEYWORDS = ["keyword1", "keyword2", "keyword3"]

**EXAMPLE OUTPUT FORMAT:**
SOFT_SKILLS = ["Communication", "Leadership", "Problem-solving"]
TECHNICAL_SKILLS = ["Python", "SQL", "Tableau"]
DOMAIN_KEYWORDS = ["Data analysis", "Business intelligence", "Machine learning"]

Text: """
{document_text.strip()}
"""
'''
        
        # Run CV extraction using unified template
        print("🔄 [CV] Starting CV skill extraction...")
        cv_prompt = create_skill_extraction_prompt("CV", cv_text)
        
        # Run JD extraction using unified template
        print("🔄 [JD] Starting JD skill extraction...")
        jd_prompt = create_skill_extraction_prompt("Job Description", jd_text)
        
        # Create unified parsing function
        def parse_skill_extraction_response(result_text: str, mode: str) -> dict:
            """Unified parsing function for both CV and JD skill extraction"""
            print(f"🔍 [{mode.upper()}] Raw response preview: {result_text[:300]}...")
            
            soft_skills, technical_skills, domain_keywords = [], [], []
            
            # Look for variable assignments like SOFT_SKILLS = ["skill1", "skill2"]
            soft_match = re.search(r'SOFT_SKILLS\s*=\s*(\[.*?\])', result_text, re.DOTALL)
            tech_match = re.search(r'TECHNICAL_SKILLS\s*=\s*(\[.*?\])', result_text, re.DOTALL)
            domain_match = re.search(r'DOMAIN_KEYWORDS\s*=\s*(\[.*?\])', result_text, re.DOTALL)
            
            if soft_match:
                try:
                    soft_skills = ast.literal_eval(soft_match.group(1))
                    print(f"🔍 [{mode.upper()}] Parsed SOFT_SKILLS variable: {len(soft_skills)} items")
                except Exception as e:
                    print(f"⚠️ [{mode.upper()}] Failed to parse SOFT_SKILLS: {e}")
                    soft_skills = []
                    
            if tech_match:
                try:
                    technical_skills = ast.literal_eval(tech_match.group(1))
                    print(f"🔍 [{mode.upper()}] Parsed TECHNICAL_SKILLS variable: {len(technical_skills)} items")
                except Exception as e:
                    print(f"⚠️ [{mode.upper()}] Failed to parse TECHNICAL_SKILLS: {e}")
                    technical_skills = []
                    
            if domain_match:
                try:
                    domain_keywords = ast.literal_eval(domain_match.group(1))
                    print(f"🔍 [{mode.upper()}] Parsed DOMAIN_KEYWORDS variable: {len(domain_keywords)} items")
                except Exception as e:
                    print(f"⚠️ [{mode.upper()}] Failed to parse DOMAIN_KEYWORDS: {e}")
                    domain_keywords = []
            
            # Fallback: try to find any Python lists in the response 
            if not soft_skills and not technical_skills and not domain_keywords:
                lists = re.findall(r'\[.*?\]', result_text, re.DOTALL)
                print(f"🔍 [{mode.upper()}] Fallback: Found {len(lists)} potential lists")
                
                if len(lists) >= 3:
                    try:
                        soft_skills = ast.literal_eval(lists[0])
                    except Exception:
                        soft_skills = []
                    try:
                        technical_skills = ast.literal_eval(lists[1])
                    except Exception:
                        technical_skills = []
                    try:
                        domain_keywords = ast.literal_eval(lists[2])
                    except Exception:
                        domain_keywords = []
            
            # If still empty, fallback to extracting from Markdown bullet points
            if not soft_skills and not technical_skills and not domain_keywords:
                def extract_bullets(section_title):
                    # Try multiple patterns for section extraction
                    patterns = [
                        rf'{section_title}:(.*?)(?:\\n## |$)',  # Original pattern
                        rf'\\*\\*{section_title}\\*\\*(.*?)(?:\\n\\*\\*|$)',  # Bold markdown
                        rf'{section_title}\\s*\\n(.*?)(?:\\n\\n|$)',  # Section with newline
                        rf'# {section_title}(.*?)(?:\\n#|$)',  # Header format
                    ]
                    
                    section_text = ''
                    for pattern in patterns:
                        match = re.search(pattern, result_text, re.DOTALL | re.IGNORECASE)
                        if match:
                            section_text = match.group(1)
                            break
                    
                    if not section_text:
                        return []
                    
                    # Extract skills from various formats
                    bullets = []
                    bullet_patterns = [
                        r'\\*\\s+(.*?)(?:\\n|$)',  # * format
                        r'-\\s+(.*?)(?:\\n|$)',   # - format
                        r'\\d+\\.\\s+(.*?)(?:\\n|$)',  # 1. format
                        r'\"([^\"]+)\"',  # Quoted items
                        r"'([^']+)'",  # Single quoted
                    ]
                    
                    for pattern in bullet_patterns:
                        found = re.findall(pattern, section_text)
                        if found:
                            bullets.extend(found)
                            break
                    
                    cleaned = []
                    for b in bullets:
                        skill = b.split(' - ')[0].strip()
                        skill = skill.strip('"').strip("'")
                        if skill and skill.lower() not in ['n/a', 'none', '']:
                            cleaned.append(skill)
                    return cleaned
                
                soft_skills = extract_bullets('SOFT SKILLS')
                technical_skills = extract_bullets('TECHNICAL SKILLS')
                domain_keywords = extract_bullets('DOMAIN KEYWORDS')
            
            # Clean skill lists
            def clean_skill_list(lst):
                return [s for s in lst if isinstance(s, str) and s.strip() and s.strip().lower() != 'n/a']
            
            soft_skills = clean_skill_list(soft_skills)
            technical_skills = clean_skill_list(technical_skills)
            domain_keywords = clean_skill_list(domain_keywords)
            
            print(f"📊 [{mode.upper()}] Final results:")
            print(f"   Soft Skills ({len(soft_skills)}): {soft_skills[:3]}{'...' if len(soft_skills) > 3 else ''}")
            print(f"   Technical Skills ({len(technical_skills)}): {technical_skills[:3]}{'...' if len(technical_skills) > 3 else ''}")
            print(f"   Domain Keywords ({len(domain_keywords)}): {domain_keywords[:3]}{'...' if len(domain_keywords) > 3 else ''}")
            
            return {
                "soft_skills": soft_skills,
                "technical_skills": technical_skills,
                "domain_keywords": domain_keywords
            }
        
        # Run both extractions in parallel
        import asyncio
        import re
        import ast
        
        async def extract_skills(prompt, mode):
            try:
                # Use the hybrid AI service to respect model selection (e.g., DeepSeek)
                full_prompt = f"You are a precise extractor of skills from professional {mode}.\n\n{prompt}"
                
                print(f"🤖 [{mode.upper()}] Using model: {hybrid_ai.provider}")
                print(f"⏱️ [{mode.upper()}] Starting API call...")
                
                # Add timeout wrapper for the API call
                try:
                    response = await asyncio.wait_for(
                        hybrid_ai.generate_response(
                            full_prompt,
                            temperature=0.1,  # Lower temperature for focused responses
                            max_tokens=1200   # Increased for comprehensive analysis with context evidence
                        ),
                        timeout=180  # 3-minute timeout per individual API call for comprehensive analysis
                    )
                    print(f"✅ [{mode.upper()}] API call completed successfully")
                except asyncio.TimeoutError:
                    print(f"⏰ [{mode.upper()}] API call timed out after 2 minutes")
                    raise Exception(f"API timeout for {mode} extraction after 2 minutes")
                
                result_text = response.strip()
                comprehensive_analysis = result_text
                
                # Use unified parsing function for both CV and JD
                parsed_skills = parse_skill_extraction_response(result_text, mode)
                
                print(f"✅ [{mode.upper()}] Extraction completed")
                return {
                    "soft_skills": parsed_skills["soft_skills"],
                    "technical_skills": parsed_skills["technical_skills"],
                    "domain_keywords": parsed_skills["domain_keywords"],
                    "comprehensive_analysis": comprehensive_analysis,
                    "raw_response": result_text
                }
            except Exception as e:
                print(f"❌ [{mode.upper()}] Extraction error: {str(e)}")
                return {"error": f"{mode.upper()} extraction error: {str(e)}"}
        
        # Run both extractions
        cv_result = await extract_skills(cv_prompt, "cv")
        jd_result = await extract_skills(jd_prompt, "jd")
        
        end_time = time.time()
        
        # Log results
        print("\n" + "="*80)
        print("🎯 PRELIMINARY ANALYSIS: Results Summary")
        print("="*80)
        print(f"CV Skills: {len(cv_result.get('technical_skills', []))} tech, {len(cv_result.get('soft_skills', []))} soft, {len(cv_result.get('domain_keywords', []))} domain")
        print(f"JD Skills: {len(jd_result.get('technical_skills', []))} tech, {len(jd_result.get('soft_skills', []))} soft, {len(jd_result.get('domain_keywords', []))} domain")
        print(f"⏱️ Total time: {end_time - start_time:.2f} seconds")
        print("="*80)
        
        # Log comprehensive analyses - extract company name from JD
        from .analysis_results_saver import AnalysisResultsSaver
        saver = AnalysisResultsSaver()
        company_name = saver.extract_company_name(jd_text)
        
        # Save company info as JSON file
        try:
            saver.save_job_info_as_json(jd_text)
            print(f"✅ [COMPANY_INFO] Company info JSON file created for: {company_name}")
        except Exception as e:
            print(f"⚠️ [COMPANY_INFO] Warning: Could not save company info JSON: {e}")

        # Save original CV text to company folder
        try:
            saver.save_original_cv_text(cv_text, company_name)
            print(f"✅ [ORIGINAL_CV] Original CV text saved for: {company_name}")
        except Exception as e:
            print(f"⚠️ [ORIGINAL_CV] Warning: Could not save original CV text: {e}")
        
        if "error" not in cv_result:
            append_output_log(cv_result["comprehensive_analysis"], company_name=company_name, tag="CV_CLAUDE_ANALYSIS")
        if "error" not in jd_result:
            append_output_log(jd_result["comprehensive_analysis"], company_name=company_name, tag="JD_CLAUDE_ANALYSIS")
        
        return {
            "cv_skills": cv_result,
            "jd_skills": jd_result,
            "total_time": end_time - start_time,
            "company_name": company_name
        }
        
    except Exception as e:
        print(f"❌ [PRELIMINARY ANALYSIS] Error: {str(e)}")
        return {"error": f"Preliminary analysis error: {str(e)}"}

@app.post("/api/set-deepseek-model")
async def set_deepseek_model(request: Request):
    """
    Quickly set the model to DeepSeek for frontend integration.
    """
    try:
        data = await request.json()
        deepseek_model = data.get('model', 'deepseek-chat')
        
        # Validate DeepSeek model
        valid_deepseek_models = ['deepseek-chat', 'deepseek-coder', 'deepseek-reasoner']
        if deepseek_model not in valid_deepseek_models:
            deepseek_model = 'deepseek-chat'  # Default to deepseek-chat
        
        # Update the global model state
        from .ai_config import model_state
        model_state.set_model(deepseek_model)
        
        print(f"🔄 [API] Switched to DeepSeek model: {deepseek_model}")
        
        return {
            "success": True,
            "message": f"Switched to {deepseek_model}",
            "current_model": model_state.get_current_model(),
            "current_provider": model_state.get_current_provider()
        }
        
    except Exception as e:
        print(f"❌ [API] Error setting DeepSeek model: {e}")
        return {
            "success": False,
            "error": str(e)
        }

@app.post("/api/update-ai-model")
async def update_ai_model(request: Request):
    """
    Update AI model configuration for a specific task.
    This endpoint allows the frontend to dynamically change AI models.
    """
    try:
        data = await request.json()
        task = data.get('task')
        model = data.get('model')
        
        if not task or not model:
            raise HTTPException(
                status_code=400, 
                detail="Both 'task' and 'model' are required"
            )
        
        # Validate task
        valid_tasks = ['DEFAULT', 'ANALYSIS', 'DETAILED_ANALYSIS', 'FAST', 'CREATIVE']
        if task not in valid_tasks:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid task. Must be one of: {valid_tasks}"
            )
        
        # Validate model - support both API models and direct model names
        from .ai_config import AI_MODELS
        valid_models = list(AI_MODELS.values())
        
        # Add DeepSeek models to valid models
        deepseek_models = ['deepseek-chat', 'deepseek-coder', 'deepseek-reasoner']
        valid_models.extend(deepseek_models)
        
        if model not in valid_models:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid model. Must be one of: {valid_models}"
            )
        
        # Update the global model state
        from .ai_config import model_state
        model_state.set_model(model)
        
        print(f"🔄 [API] Updated {task} model to: {model}")
        
        return {
            "success": True,
            "message": f"Updated {task} model to {model}",
            "task": task,
            "model": model
        }
        
    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [API] Error updating AI model: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to update AI model: {str(e)}"
        )

# Job Queue API Endpoints
@app.post("/api/jobs/start")
async def start_job(request: Request):
    """Start a background job"""
    try:
        data = await request.json()
        job_type = data.get('job_type')
        job_data = data.get('job_data', {})
        
        if not job_type:
            raise HTTPException(status_code=400, detail="job_type is required")
        
        job_id = job_queue.submit_job(job_type, job_data)
        
        return {
            "job_id": job_id,
            "status": "pending",
            "message": f"Job {job_id} started"
        }
        
    except Exception as e:
        print(f"❌ [JOB_API] Error starting job: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start job: {str(e)}"
        )

@app.get("/api/jobs/{job_id}/status")
async def get_job_status(job_id: str):
    """Get the status of a background job"""
    try:
        job_status = job_queue.get_job_status(job_id)
        
        if not job_status:
            raise HTTPException(status_code=404, detail="Job not found")
        
        return job_status
        
    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [JOB_API] Error getting job status: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get job status: {str(e)}"
        )

@app.post("/api/jobs/{job_id}/cancel")
async def cancel_job(job_id: str):
    """Cancel a background job"""
    try:
        success = job_queue.cancel_job(job_id)
        
        if not success:
            raise HTTPException(
                status_code=400,
                detail="Job cannot be cancelled (not found or already completed)"
            )
        
        return {
            "success": True,
            "message": f"Job {job_id} cancelled"
        }
        
    except HTTPException as http_e:
        raise http_e
    except Exception as e:
        print(f"❌ [JOB_API] Error cancelling job: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to cancel job: {str(e)}"
        )

@app.get("/api/jobs/queue/stats")
async def get_queue_stats():
    """Get job queue statistics"""
    try:
        stats = job_queue.get_queue_stats()
        return stats
        
    except Exception as e:
        print(f"❌ [JOB_API] Error getting queue stats: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get queue stats: {str(e)}"
        )

