from src.generate_tailored_cv import generate_tailored_cv
from pydantic import BaseModel
from docx import Document
import os

class TestRequest(BaseModel):
    cv_filename: str = 'test.docx'
    jd_text: str = 'Looking for a Software Engineer at Google. Must have Python and ML experience.'
    custom_prompt: str = 'Tailor CV for software role'
    source: str = 'test'
    use_last_tested: bool = False
    is_new_cv: bool = False
    job_link: str = 'https://test-job-link.com'

def create_test_cv():
    print("\nCreating test CV file...")
    # Create uploads directory if it doesn't exist
    os.makedirs('uploads', exist_ok=True)
    
    # Create a simple test CV
    doc = Document()
    doc.add_heading('Test CV', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('• Python\n• Machine Learning\n• Software Development')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Software Engineer at Tech Corp')
    
    # Save the test CV
    cv_path = os.path.join('uploads', 'test.docx')
    doc.save(cv_path)
    print(f"✅ Test CV created at: {cv_path}")

def test_cv_generation():
    print("\n=== Testing CV Generation ===")
    
    # Create test CV first
    create_test_cv()
    
    print("\nCreating test request...")
    request = TestRequest()
    
    print("\nCalling generate_tailored_cv...")
    try:
        result = generate_tailored_cv(request)
        print("\n=== Generation Result ===")
        print(f"Result: {result}")
    except Exception as e:
        print(f"\n❌ Error during generation: {str(e)}")

if __name__ == "__main__":
    test_cv_generation() 